#!/usr/bin/env python3
"""Resume Tailor Web App — Flask backend."""

import io
import json
import os
import re
import traceback
import unicodedata
from datetime import datetime
from urllib.parse import urlparse

import requests
from flask import Flask, render_template, request, jsonify, send_file
from fpdf import FPDF

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB max upload


# ─────────────────────────────────────────────
# Resume file parsing
# ─────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Iterate body elements in document order (paragraphs AND tables)
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            para = Paragraph(element, doc)
            if para.text.strip():
                parts.append(para.text.strip())
        elif tag == "tbl":
            table = Table(element, doc)
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

    return "\n".join(parts)


def parse_pdf(file_bytes: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text_parts.append(t)
    return "\n".join(text_parts)


def parse_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


PARSERS = {
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".txt": parse_txt,
}


def _extract_paragraph_text_with_links(paragraph) -> str:
    """Extract paragraph text, replacing field-code hyperlink display text with actual URLs."""
    from docx.oxml.ns import qn

    # Collect field-code hyperlinks: HYPERLINK "url" -> display text
    # Field codes use: fldChar(begin) -> instrText(HYPERLINK "url") -> fldChar(separate) -> runs(display) -> fldChar(end)
    runs = paragraph._element.findall(qn("w:r"))
    parts = []
    in_field = False
    hyperlink_url = None
    collecting_display = False
    display_parts = []

    for run in runs:
        fld_char = run.find(qn("w:fldChar"))
        instr_text = run.find(qn("w:instrText"))
        text_el = run.find(qn("w:t"))

        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_field = True
                hyperlink_url = None
                display_parts = []
            elif fld_type == "separate":
                collecting_display = True
            elif fld_type == "end":
                # Emit the collected hyperlink
                if hyperlink_url and display_parts:
                    display = "".join(display_parts).strip()
                    url = hyperlink_url.replace("mailto:", "")
                    parts.append(f"[{display}]({url})")
                in_field = False
                collecting_display = False
                hyperlink_url = None
                display_parts = []
        elif instr_text is not None and in_field:
            m = re.search(r'HYPERLINK\s+"([^"]+)"', instr_text.text or "")
            if m:
                hyperlink_url = m.group(1)
        elif text_el is not None:
            if collecting_display and hyperlink_url:
                display_parts.append(text_el.text or "")
            elif not in_field:
                parts.append(text_el.text or "")

    result = "".join(parts).strip()
    return result if result else paragraph.text.strip()


def extract_personal_info_docx(file_bytes: bytes) -> dict:
    """Extract name, contact details from DOCX tables/headers before AI touches it."""
    from docx import Document
    from docx.table import Table

    doc = Document(io.BytesIO(file_bytes))
    info = {"name": "", "contact": ""}

    # Check tables first (common resume format: header table with name + contact)
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "tbl":
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        style = p.style.name if p.style else ""
                        text = _extract_paragraph_text_with_links(p)
                        if not text:
                            continue
                        # Title style = candidate name
                        if style == "Title" and not info["name"]:
                            info["name"] = text
                        # Build contact from non-title, non-subtitle paragraphs
                        elif style not in ("Title", "Subtitle") and text:
                            if info["contact"]:
                                info["contact"] += " | " + text
                            else:
                                info["contact"] = text
            # Only check the first table (header table)
            break

    # Clean up contact: collapse whitespace, normalize separators
    if info["contact"]:
        info["contact"] = re.sub(r"\s+", " ", info["contact"]).strip()

    return info


def extract_personal_info_text(resume_text: str) -> dict:
    """Extract name and contact details from plain text resume (works for PDF/TXT).

    Looks at the first few non-empty lines for name, email, phone, URLs.
    """
    lines = resume_text.splitlines()
    non_empty = [l.strip() for l in lines if l.strip()][:8]

    info = {"name": "", "contact": ""}
    contact_parts = []

    for i, line in enumerate(non_empty):
        # First non-empty line is usually the name
        if i == 0:
            # Skip if it looks like contact info
            if "@" not in line and "http" not in line.lower() and not re.search(r"\(\d{3}\)", line):
                info["name"] = line.strip()
                continue

        # Collect contact-like lines (contain email, phone, URL, location with pipe separators)
        if any(indicator in line.lower() for indicator in ["@", "http", "linkedin", "github", "|"]) or re.search(r"\(\d{3}\)", line):
            contact_parts.append(line.strip())
        elif i <= 2 and re.match(r"^[A-Za-z\s,]+,\s*[A-Z]{2}", line):
            # Looks like "City, ST" line
            contact_parts.append(line.strip())

    if contact_parts:
        # Join and clean up PDF extraction artifacts (extra spaces around hyphens/dashes)
        raw_contact = " | ".join(contact_parts)
        # Fix "300 -8788" → "300-8788", "ryan -murphy" → "ryan-murphy"
        raw_contact = re.sub(r"\s+-\s*", "-", raw_contact)
        raw_contact = re.sub(r"\s*-\s+", "-", raw_contact)
        # Fix spaces before punctuation: "Beach , FL" → "Beach, FL"
        raw_contact = re.sub(r"\s+,", ",", raw_contact)
        # Collapse multiple spaces
        raw_contact = re.sub(r"\s{2,}", " ", raw_contact)
        # Collapse multiple pipes
        raw_contact = re.sub(r"\|\s*\|", "|", raw_contact)
        info["contact"] = raw_contact.strip()

    # Also clean name
    if info["name"]:
        info["name"] = re.sub(r"\s{2,}", " ", info["name"]).strip()

    return info


def _extract_original_title(resume_text: str) -> str:
    """Extract the candidate's job title from the resume text.

    Looks at the first ~10 non-empty lines for a title-like line
    (e.g. "Senior Software Engineer", "Full Stack Developer").
    """
    title_keywords = [
        "engineer", "developer", "architect", "manager", "designer",
        "analyst", "consultant", "scientist", "administrator", "lead",
        "director", "specialist", "coordinator", "intern", "associate",
        "technician", "strategist", "officer", "president", "founder",
    ]
    lines = resume_text.splitlines()
    non_empty = [l.strip() for l in lines if l.strip()][:10]

    for line in non_empty:
        lower = line.lower()
        # Skip lines that look like contact info, section headers, or name (usually first line)
        if "@" in line or "|" in line or "http" in lower:
            continue
        if any(kw in lower for kw in title_keywords):
            # Likely the title line
            return line.strip()

    return ""


# ─────────────────────────────────────────────
# Default tailoring prompt
# ─────────────────────────────────────────────

DEFAULT_TAILORING_PROMPT = """
================================================================================
MASTER PROMPT: RESUME TAILORING ENGINE (v1.0)
================================================================================

You are a world-class Resume Strategist and ATS Optimization Expert with 15+ years
of experience in technical recruiting, hiring pipeline optimization, and career
coaching for senior software engineers. You have deep expertise in how Applicant
Tracking Systems (Greenhouse, Lever, Workday, iCIMS, Taleo) parse, score, and rank
resumes. You understand exactly how hiring managers, technical recruiters, and
engineering directors evaluate resumes for senior/staff-level engineering roles.

================================================================================
YOUR MISSION
================================================================================

You will receive TWO inputs:
  1. A CANDIDATE RESUME (the source of truth for experience, skills, and facts)
  2. A JOB DESCRIPTION (the target you are optimizing toward)

Your task is to produce a TAILORED RESUME that maximizes the candidate's chances of:
  (a) Passing ATS keyword filters and scoring algorithms
  (b) Getting shortlisted by a technical recruiter (6-second scan test)
  (c) Impressing the hiring manager during a detailed read
  (d) Setting up strong talking points for interviews

================================================================================
PHASE 1: DEEP JD ANALYSIS (Do this FIRST before touching the resume)
================================================================================

Before making ANY changes to the resume, perform this comprehensive JD analysis.
Output your analysis so the reasoning is transparent.

1.1 — ROLE IDENTITY EXTRACTION
    - What is the EXACT title? (Use this to retitle the resume)
    - What is the seniority level? (Junior / Mid / Senior / Staff / Principal)
    - What is the role TYPE? (IC, Tech Lead, Architect, Manager-of-one, Player-Coach)
    - Is it frontend-heavy, backend-heavy, full-stack equal, or has a PRIMARY focus?
    - What DOMAIN is the company in? (Healthcare, Finance, E-commerce, SaaS, etc.)
    - Remote / Hybrid / On-site?

1.2 — HARD REQUIREMENTS EXTRACTION (Must-Have)
    Extract EVERY hard requirement. Categorize them:

    [LANGUAGES]       — e.g., TypeScript, C#, Python
    [FRAMEWORKS]      — e.g., React, Angular, .NET, Next.js
    [ARCHITECTURE]    — e.g., Microservices, Event-Driven, Distributed Systems
    [CLOUD/DEVOPS]    — e.g., AWS, Azure, Docker, Kubernetes, CI/CD
    [DATA]            — e.g., PostgreSQL, MongoDB, SQL Server
    [AI/ML]           — e.g., LLM, RAG, Inference Pipelines, Model Deployment
    [SECURITY]        — e.g., Secure Coding, DevSecOps, Compliance
    [SOFT SKILLS]     — e.g., Mentoring, Cross-functional Collaboration
    [YEARS]           — Minimum years of experience required
    [EDUCATION]       — Degree requirements if any

1.3 — PREFERRED/BONUS QUALIFICATIONS
    Extract ALL preferred qualifications separately. These are differentiators.
    Mark each as [HIGH-VALUE BONUS] or [NICE-TO-HAVE].

1.4 — HIDDEN KEYWORDS & IMPLICIT REQUIREMENTS
    Identify keywords that are IMPLIED but not explicitly stated:
    - If JD says "scalable frontend architectures" -> implies: Component Libraries,
      Design Systems, State Management, Performance Budgets
    - If JD says "production ownership" -> implies: On-call, Incident Response,
      Runbooks, SLA Management, Post-mortems
    - If JD says "AI integration" -> implies: Prompt Engineering, Token Optimization,
      Model Evaluation, A/B Testing of AI features
    - If JD says "observability" -> implies: DataDog/Grafana/NewRelic, Distributed
      Tracing, Log Aggregation, Alerting Thresholds

1.5 — TONE & CULTURE SIGNALS
    Analyze the JD's language for cultural signals:
    - Does it say "ownership" / "accountability" -> Emphasize autonomous leadership
    - Does it say "collaborate" / "cross-functional" -> Emphasize teamwork
    - Does it say "systems thinking" -> Emphasize architectural judgment
    - Does it say "measurable impact" -> Every bullet MUST have metrics
    - Does it say "not a feature-delivery role" -> De-emphasize "built features",
      emphasize "owned systems", "designed architecture", "improved reliability"

1.6 — PRIORITY STACK RANKING
    Rank all extracted requirements by importance (based on JD ordering, repetition,
    and emphasis). The resume should mirror this priority order:

    Priority 1 (MUST dominate resume): _______________
    Priority 2 (MUST be prominent):    _______________
    Priority 3 (Should appear clearly): _______________
    Priority 4 (Should be mentioned):   _______________
    Priority 5 (Nice to include):       _______________

================================================================================
PHASE 2: RESUME AUDIT (Map existing resume against JD)
================================================================================

2.1 — KEYWORD MATCH MATRIX
    Create a matrix:
    | JD Requirement          | Present in Resume? | Where?          | Strength |
    |-------------------------|--------------------|-----------------|----------|
    | TypeScript              | YES                | Summary, Skills | Strong   |
    | Core Web Vitals         | NO                 | -               | Missing  |
    | Event-Driven APIs       | Partial            | Microsoft only  | Weak     |

    Mark each as: STRONG / ADEQUATE / WEAK / MISSING

2.2 — EXPERIENCE ALIGNMENT SCORING
    For each work experience entry, score alignment to JD (0-100%):
    - Which bullets already align?
    - Which bullets are irrelevant to this JD?
    - Which bullets can be REFRAMED to align?
    - What MISSING experiences from this role should be surfaced?

2.3 — GAP IDENTIFICATION
    List ALL gaps between JD and resume:
    [CRITICAL GAPS]   — Required skills/keywords completely missing
    [LANGUAGE GAPS]    — Skill exists but wrong terminology used
    [EMPHASIS GAPS]    — Skill exists but buried or underweighted
    [STRUCTURAL GAPS]  — Information in wrong section or wrong order

================================================================================
PHASE 3: RESUME REWRITING RULES
================================================================================

FOLLOW THESE RULES WITH ZERO EXCEPTIONS:

3.1 — GOLDEN RULE: NEVER FABRICATE
    - NEVER invent experience, skills, projects, metrics, or companies
    - NEVER inflate numbers beyond what's stated in the original resume
    - NEVER add technologies the candidate hasn't used
    - You may REFRAME, REWORD, and REORDER — but NEVER FABRICATE
    - If a JD requirement has NO match in the resume, leave it out — do NOT fake it

3.2 — TITLE
    - ALWAYS keep the candidate's EXACT original title from the resume — no exceptions
    - Do NOT change, adapt, or customize the title based on the JD
    - The title must be copied verbatim from the original resume

3.3 — SUMMARY REWRITING (Most critical section — recruiters read this first)

    Structure (4-6 sentences max):
    Sentence 1: [Seniority] + [Primary Identity from JD] + [Years] + [Domain Match]
    Sentence 2: [PRIMARY technical strength matching JD Priority 1]
    Sentence 3: [SECONDARY technical strength matching JD Priority 2]
    Sentence 4: [AI/Cloud/Specialized skill matching JD Priority 3]
    Sentence 5: [Ownership/Leadership/Impact signal matching JD tone]
    Sentence 6: [Differentiator — regulated industry, scale, mentoring]

    Rules:
    - First 15 words must contain the JD's top 2-3 keywords
    - Include EXACT phrases from JD where natural (ATS matching)
    - Quantify at least 2 claims (years, scale, percentage improvements)
    - Mirror the JD's tone (ownership-focused vs. collaboration-focused)
    - Do NOT use first person ("I built...") — use telegraphic style
    - Do NOT use buzzwords without substance ("passionate", "rockstar", "guru")

3.4 — SKILLS SECTION REWRITING

    Rules:
    - REORDER skill categories to match JD priority stack (Phase 1.6)
    - RENAME skill categories to mirror JD section headers EXACTLY
      (If JD says "Frontend Architecture & Performance Leadership",
       the skill category should be "Frontend Architecture & Performance")
    - Within each category, list JD-mentioned skills FIRST, then supporting skills
    - ADD skills that exist in the resume but are scattered in experience bullets
    - REMOVE or DEPRIORITIZE skills not relevant to this specific JD
    - Use the EXACT terminology from the JD:
      JD says "C#" -> use "C#" (not "C Sharp")
      JD says "Kubernetes" -> use "Kubernetes" (not "K8s")
      JD says "event-driven" -> use "Event-Driven" (not "message-based")
    - Include IMPLICIT skills from Phase 1.4 if candidate has them

3.5 — EXPERIENCE BULLET REWRITING (The most impactful section)

    For EACH bullet point, apply this framework:

    FORMULA: [POWER VERB] + [WHAT you did using JD KEYWORDS] + [SCALE/SCOPE] +
             [MEASURABLE RESULT with METRIC]

    Example transformation:
    BEFORE: "Built React components for dashboards"
    AFTER:  "Architected scalable React and TypeScript frontend systems for clinical
             dashboards serving 10k+ daily users, optimizing Core Web Vitals and
             reducing Time-to-Interactive by 40%"

    Power Verb Selection (match to JD tone):
    - For architecture roles: Architected, Designed, Engineered, Established
    - For ownership roles:    Owned, Led, Drove, Spearheaded
    - For optimization roles: Optimized, Reduced, Improved, Accelerated
    - For AI/innovation:      Implemented, Integrated, Developed, Productionized
    - For mentoring:          Mentored, Guided, Established, Influenced

    AVOID these weak verbs: Helped, Assisted, Participated, Worked on, Was responsible

    Rules for each bullet:
    a) Start with the STRONGEST action verb (no two bullets should start with same verb)
    b) Front-load JD keywords in the first half of the bullet
    c) Include at least ONE metric per bullet (%, #, scale, time saved)
    d) Connect the action to BUSINESS IMPACT (not just technical output)
    e) Use JD's EXACT phrases where natural
    f) If a bullet doesn't serve this specific JD, either REFRAME it or REMOVE it
    g) Aim for 5-7 bullets per recent role, 4-5 for older roles
    h) Most impactful/JD-aligned bullets go FIRST within each role

3.6 — EXPERIENCE ORDERING WITHIN EACH ROLE
    Reorder bullets within each job to match JD priority:
    1st bullet: PRIMARY JD focus (e.g., Frontend Architecture)
    2nd bullet: SECONDARY JD focus (e.g., Backend/Systems)
    3rd bullet: AI/ML integration (if applicable)
    4th bullet: Cloud/DevOps/Observability
    5th bullet: Data/Security
    6th bullet: Leadership/Mentoring
    7th bullet: Other measurable impact

3.7 — COMPANY CONTEXT ANNOTATIONS
    For each role, add a brief context tag if the company domain matches JD preferences:
    - If JD prefers regulated industries -> tag "(Healthcare - Regulated Industry)"
    - If JD prefers high-scale -> tag "(E-Commerce - 2M+ Monthly Users)"
    - This signals domain alignment immediately to the recruiter

3.8 — EDUCATION & CERTIFICATIONS
    - Keep education section brief
    - If candidate has relevant certifications (AWS, Azure, etc.), ensure they're visible
    - If JD mentions specific certifications, and candidate has them, HIGHLIGHT them

================================================================================
PHASE 4: ATS OPTIMIZATION CHECKLIST
================================================================================

After rewriting, verify ALL of the following:

4.1 — KEYWORD DENSITY
    [ ] Every REQUIRED skill from JD appears at least 2x in resume
        (once in Skills, once in Experience)
    [ ] Every PREFERRED skill appears at least 1x
    [ ] Job title or close variant appears in resume title AND summary
    [ ] Industry-specific terms from JD are present

4.2 — FORMATTING (ATS-Safe)
    [ ] No tables, columns, graphics, or images (ATS can't parse them)
    [ ] No headers/footers with critical info (ATS often skips them)
    [ ] Standard section headers: SUMMARY, SKILLS, EXPERIENCE, EDUCATION
    [ ] Dates in consistent format (MM/YYYY or Month YYYY)
    [ ] No special characters that might break parsing (em-dashes, smart quotes)
    [ ] Plain text or simple formatting only
    [ ] Contact info at the very top (not in a sidebar)

4.3 — LENGTH & DENSITY
    [ ] The resume should be as long as the original — preserve ALL content and bullets
    [ ] No single bullet exceeds 3 lines
    [ ] White space is adequate for readability
    [ ] Most recent 2 roles get the most space (60% of experience section)

4.4 — CONSISTENCY
    [ ] Same tense throughout (past tense for previous roles, present for current)
    [ ] Consistent bullet formatting (all start with verb, all have metrics)
    [ ] Technology names spelled identically throughout (TypeScript, not TS sometimes)
    [ ] No orphan skills (skills in Skills section should appear in Experience too)

================================================================================
PHASE 5: FINAL QUALITY CHECKS
================================================================================

Before delivering, verify:

5.1 — THE 6-SECOND TEST
    Read ONLY the title, summary, and skill headers. In 6 seconds, is it clear that
    this candidate matches the JD? If not, rewrite.

5.2 — THE KEYWORD MATCH TEST (HARD REQUIREMENT — NOT OPTIONAL)
    TARGET: ATS keyword-coverage score >= 85% on the FIRST PASS.

    Before you finalize the JSON, perform this exact algorithm in your head:
    a) Extract every distinct technical noun, framework, tool, methodology, and
       domain term from the JD (call this set K). Treat multi-word terms
       ("event-driven architecture") as single keywords.
    b) For each k in K, check whether k (or an obvious surface form: case-insensitive,
       hyphen/space-insensitive, common synonym already established as equivalent)
       appears at least ONCE in the candidate's tailored resume.
    c) Compute coverage = |present| / |K|. This MUST be >= 0.85.
    d) If coverage < 0.85, REVISE: surface keywords from the resume into Skills
       and Experience bullets — but ONLY if the candidate genuinely has that
       experience. Never fabricate. If a keyword cannot be sourced from the
       original resume, leave it out.
    e) Re-compute coverage. Iterate until >= 0.85 OR until you have exhausted
       every legitimate keyword the resume can support.

    REQUIRED SURFACE-AREA RULES (these dramatically increase ATS scores):
    - Every required JD keyword you DO have appears at least 2x: once in Skills,
      once in an Experience bullet (or Summary).
    - Job-title synonyms from the JD appear in Summary's first sentence.
    - The exact JD section labels you mirror (e.g., "Frontend Architecture &
      Performance") appear as Skills category names.
    - Acronyms appear in BOTH expanded and short form ("Continuous Integration (CI)")
      so ATS keyword matchers hit on either form.

5.3 — THE AUTHENTICITY TEST
    Compare every claim in the tailored resume against the original resume.
    Flag ANYTHING that could be seen as exaggerated or fabricated.

5.4 — THE "SO WHAT" TEST
    Every bullet should answer: "So what? Why does this matter to THIS employer?"
    If a bullet doesn't connect to JD priorities, cut it or reframe it.

5.5 — THE SPECIFICITY TEST
    Remove any vague language:
    - "various technologies" -> name the specific technologies
    - "improved performance" -> "reduced API latency by 25%"
    - "worked with teams" -> "mentored 5 engineers across 3 teams"
    - "large-scale systems" -> "systems handling 5M+ records/month"

5.6 — THE RECENCY BIAS CHECK
    Most recent role should have the MOST bullets and STRONGEST alignment.
    Older roles can be condensed — but don't lose critical JD-matching experience.

================================================================================
CONSTRAINTS & ANTI-PATTERNS
================================================================================

NEVER DO THESE:
  x  Do NOT copy-paste JD language into resume without adapting it naturally
  x  Do NOT use the same metric in multiple bullets
  x  Do NOT start consecutive bullets with the same word
  x  Do NOT include technologies the candidate hasn't used (even if JD requires them)
  x  Do NOT remove ALL non-JD experience (shows breadth and avoids looking "too perfect")
  x  Do NOT use subjective self-assessments ("excellent communicator", "strong leader")
  x  Do NOT include objective statements ("Seeking a role in...")
  x  Do NOT include references or "References available upon request"
  x  Do NOT use acronyms without spelling them out at least once
  x  Do NOT sacrifice readability for keyword stuffing
  x  The resume can be as long as needed — do NOT cut content to fit a page limit

================================================================================
NOW EXECUTE
================================================================================

Apply ALL phases above to the following inputs:
""".strip()


# ─────────────────────────────────────────────
# Resume section detection
# ─────────────────────────────────────────────

# Common resume section headings (case-insensitive matching)
_KNOWN_SECTIONS = [
    "summary", "professional summary", "profile", "objective",
    "skills", "technical skills", "core competencies", "technologies",
    "experience", "professional experience", "work experience", "employment",
    "career experience",
    "education", "academic background",
    "projects", "personal projects", "key projects",
    "certifications", "certificates", "licenses",
    "awards", "honors", "achievements",
    "publications", "research",
    "volunteer", "volunteer experience", "community involvement",
    "languages", "interests", "hobbies",
    "references",
]

def detect_resume_sections(resume_text: str) -> list[str]:
    """Detect section headings present in the resume text.

    Returns a list of section names in the order they appear.
    Skips the first few non-empty lines (name/contact area) to avoid false positives.
    """
    lines = resume_text.splitlines()
    found = []
    seen = set()

    # Skip the first few non-empty lines (typically name, title, contact info)
    non_empty_count = 0
    skip_until = 0
    for i, line in enumerate(lines):
        if line.strip():
            non_empty_count += 1
            if non_empty_count >= 4:
                skip_until = i
                break

    for line in lines[skip_until:]:
        stripped = line.strip().rstrip(":").strip()
        if not stripped or len(stripped) > 60:
            continue
        # Normalize: collapse multiple spaces, lowercase
        normalized = re.sub(r"\s+", " ", stripped).lower()

        # Match against known section headings only
        for known in _KNOWN_SECTIONS:
            if normalized == known:
                # Use the known canonical form for consistency
                canonical = known.title()
                if known not in seen:
                    seen.add(known)
                    found.append(canonical)
                break

    # Fallback: if nothing detected, return a standard set
    if not found:
        found = ["Summary", "Skills", "Experience", "Education"]

    # Check for implicit summary: if there's a long paragraph between
    # the header area and the first detected section, add "Summary" at the start
    if found and not any(s.lower() in ("summary", "professional summary", "profile", "objective") for s in found):
        first_section = found[0].lower()
        # Look for a long text block before the first section heading
        for line in lines[skip_until:]:
            stripped = line.strip()
            normalized = re.sub(r"\s+", " ", stripped).lower().rstrip(":")
            if normalized == first_section:
                break
            # If we find a line longer than 80 chars, it's likely a summary paragraph
            if len(stripped) > 80:
                found.insert(0, "Summary")
                break

    return found


def _classify_section(name: str) -> str:
    """Classify a section name into a rendering type."""
    lower = name.lower()
    if any(k in lower for k in ["summary", "profile", "objective"]):
        return "text"
    if any(k in lower for k in ["skill", "competenc", "technolog"]):
        return "skills"
    if any(k in lower for k in ["experience", "employment", "work", "career"]):
        return "experience"
    if any(k in lower for k in ["education", "academic"]):
        return "education"
    if any(k in lower for k in ["project"]):
        return "projects"
    if any(k in lower for k in ["certification", "certificate", "license"]):
        return "simple_list"
    if any(k in lower for k in ["award", "honor", "achievement"]):
        return "simple_list"
    if any(k in lower for k in ["publication", "research"]):
        return "simple_list"
    if any(k in lower for k in ["volunteer", "community"]):
        return "experience"
    if any(k in lower for k in ["language", "interest", "hobb"]):
        return "simple_list"
    return "simple_list"


# ─────────────────────────────────────────────
# AI API calls (Claude + Gemini)
# ─────────────────────────────────────────────

def _build_tailor_user_msg(resume_text: str, prompt_text: str, jd_text: str, detected_sections: list[str] | None = None) -> str:
    """Build the full prompt for resume tailoring (shared across providers)."""
    system_msg = prompt_text.strip()

    # Build dynamic section schema based on detected sections
    if not detected_sections:
        detected_sections = ["Summary", "Skills", "Experience", "Education"]

    section_examples = []
    for sec_name in detected_sections:
        sec_type = _classify_section(sec_name)
        if sec_type == "text":
            section_examples.append(
                f'    {{"type": "text", "heading": "{sec_name}", "content": "The full {sec_name.lower()} paragraph"}}'
            )
        elif sec_type == "skills":
            section_examples.append(
                f'    {{"type": "skills", "heading": "{sec_name}", "items": [{{"category": "Category Name", "items": "skill1, skill2, skill3"}}]}}'
            )
        elif sec_type == "experience":
            section_examples.append(
                f'    {{"type": "experience", "heading": "{sec_name}", "items": [{{"company": "Company Name", "job_title": "Title", "dates": "MM/YYYY - MM/YYYY", "bullets": ["bullet 1", "bullet 2"]}}]}}'
            )
        elif sec_type == "education":
            section_examples.append(
                f'    {{"type": "education", "heading": "{sec_name}", "items": [{{"degree": "Degree Name", "school": "School Name", "dates": "MM/YYYY - MM/YYYY", "location": "City, ST"}}]}}'
            )
        elif sec_type == "projects":
            section_examples.append(
                f'    {{"type": "projects", "heading": "{sec_name}", "items": [{{"name": "Project Name", "description": "Brief description", "technologies": "tech1, tech2", "bullets": ["detail 1", "detail 2"]}}]}}'
            )
        else:  # simple_list
            section_examples.append(
                f'    {{"type": "simple_list", "heading": "{sec_name}", "items": ["Item 1", "Item 2"]}}'
            )

    sections_json = ",\n".join(section_examples)
    sections_list = ", ".join(detected_sections)

    user_msg = f"""Here are the two inputs:

=== CANDIDATE RESUME ===
{resume_text}

=== JOB DESCRIPTION ===
{jd_text}

=== INSTRUCTIONS ===
Apply every phase from your system prompt to these inputs.

CRITICAL RULES:
1. Do NOT change the candidate's title. Keep the EXACT title from the original resume.
2. The output MUST have the SAME sections as the original resume, in the SAME order. The detected sections are: [{sections_list}].
3. Do NOT drop, merge, or add sections. Mirror the original resume's structure exactly.

IMPORTANT: Return ONLY the tailored resume as a JSON object with this exact structure (no change log, no interview prep — just the resume):

```json
{{
  "name": "Candidate Name",
  "title": "EXACT title from original resume — do NOT change this",
  "contact": "Location | email",
  "sections": [
{sections_json}
  ]
}}
```

Return ONLY the JSON object, no markdown code fences, no extra text. Just pure JSON.
"""
    return f"{system_msg}\n\n{user_msg}"


def _strip_code_fences(raw: str) -> str:
    """Strip markdown code fences from AI response."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    return raw


def _repair_json(raw: str) -> str:
    """Best-effort JSON repair for common LLM output issues.

    Fixes:
    - Literal newlines / tabs / carriage returns inside string values
      (Gemini json-mode sometimes emits unescaped these in long bullets).
    - Trailing commas before } or ].
    - Unterminated trailing string (closes with a single `"` then balances brackets).
    """
    # Walk the string with a tiny state machine, escaping control chars inside strings.
    out = []
    in_string = False
    escape = False
    for ch in raw:
        if escape:
            out.append(ch)
            escape = False
            continue
        if ch == "\\":
            out.append(ch)
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            out.append(ch)
            continue
        if in_string:
            if ch == "\n":
                out.append("\\n")
                continue
            if ch == "\r":
                out.append("\\r")
                continue
            if ch == "\t":
                out.append("\\t")
                continue
        out.append(ch)
    repaired = "".join(out)

    # Strip trailing commas before } or ]
    repaired = re.sub(r",(\s*[}\]])", r"\1", repaired)

    # Try parsing as-is; if it still has an unterminated string at the end,
    # close it and balance brackets.
    try:
        json.loads(repaired)
        return repaired
    except json.JSONDecodeError:
        pass

    # If we ended in the middle of a string, close it.
    if in_string:
        repaired += '"'
    # Balance brackets/braces by counting and appending what's missing.
    open_curly = repaired.count("{") - repaired.count("}")
    open_square = repaired.count("[") - repaired.count("]")
    # Trim a trailing comma that may now sit right before the synthetic closer
    repaired = re.sub(r",\s*$", "", repaired)
    repaired += "]" * max(0, open_square)
    repaired += "}" * max(0, open_curly)
    return repaired


def _safe_json_loads(raw: str):
    """Try json.loads; on failure, run the lightweight repair; if that still fails,
    fall back to the heavy-duty `json-repair` library which handles unquoted keys,
    single-quoted strings, missing commas, Python-style values, and other LLM quirks."""
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    try:
        return json.loads(_repair_json(raw))
    except json.JSONDecodeError:
        pass
    # Last resort: json-repair. It returns either a parsed object or the repaired string.
    from json_repair import repair_json
    repaired = repair_json(raw, return_objects=True)
    if repaired in ("", None):
        # Library returns "" when it gives up.
        raise json.JSONDecodeError("json-repair could not parse the response.", raw, 0)
    return repaired


# ─────────────────────────────────────────────
# AI Provider registry (multi-provider support)
# ─────────────────────────────────────────────

# Default model per provider. Frontend can override via `model` form field.
PROVIDER_DEFAULTS = {
    "anthropic": "claude-sonnet-4-5-20250929",
    # gemini-2.5-flash-lite has a 1,500/day free-tier quota (vs. 20/day for flash).
    # The ATS pipeline does 4-6 calls per tailor, so flash burns the daily budget in
    # 3-4 runs. flash-lite is plenty fast and survives heavier use on the free tier.
    "gemini":    "gemini-2.5-flash-lite",
    "openai":    "gpt-4o-mini",
    "openrouter": "openrouter/auto",
    "groq":      "llama-3.3-70b-versatile",
    "together":  "meta-llama/Llama-3.3-70B-Instruct-Turbo",
    "openai_compatible": "",   # user must supply
    "ollama":    "llama3.1",
}

# OpenAI-compatible providers — share the same code path, just differ in base_url.
OPENAI_COMPATIBLE_BASE_URLS = {
    "openai":      "https://api.openai.com/v1",
    "openrouter":  "https://openrouter.ai/api/v1",
    "groq":        "https://api.groq.com/openai/v1",
    "together":    "https://api.together.xyz/v1",
    "ollama":      "http://localhost:11434/v1",   # Ollama's OpenAI shim
    # "openai_compatible" — base_url comes from the request
}


def _split_system_user(prompt: str) -> tuple[str, str]:
    """The legacy prompt builder concatenates `system\n\nuser`. Split it back."""
    if "\n\n" in prompt:
        sys_part, user_part = prompt.split("\n\n", 1)
        return sys_part, user_part
    return "", prompt


def _call_anthropic(api_key: str, model: str, prompt: str, max_tokens: int, json_mode: bool, base_url: str | None) -> str:
    import anthropic
    client = anthropic.Anthropic(api_key=api_key, base_url=base_url) if base_url else anthropic.Anthropic(api_key=api_key)
    system_text, user_text = _split_system_user(prompt)
    # Cache the (large) system prompt so repeat tailorings cost ~10% of normal.
    system_blocks = [{"type": "text", "text": system_text, "cache_control": {"type": "ephemeral"}}] if system_text else []
    kwargs = {
        "model": model,
        "max_tokens": max_tokens,
        "temperature": 0,
        "messages": [{"role": "user", "content": user_text}],
    }
    if system_blocks:
        kwargs["system"] = system_blocks
    message = client.messages.create(**kwargs)
    return message.content[0].text.strip()


def _call_gemini(api_key: str, model: str, prompt: str, max_tokens: int, json_mode: bool, base_url: str | None) -> str:
    from google import genai
    from google.genai import types
    client = genai.Client(api_key=api_key)
    config_kwargs = {"max_output_tokens": max_tokens, "temperature": 0}
    if json_mode:
        config_kwargs["response_mime_type"] = "application/json"
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=types.GenerateContentConfig(**config_kwargs),
    )
    return (response.text or "").strip()


def _call_openai_compatible(api_key: str, model: str, prompt: str, max_tokens: int, json_mode: bool, base_url: str) -> str:
    """Works for OpenAI, OpenRouter, Groq, Together, Ollama, LM Studio, any OpenAI-compatible endpoint."""
    from openai import OpenAI
    # Ollama's OpenAI shim accepts any string as api_key; pass a placeholder if missing.
    client = OpenAI(api_key=api_key or "ollama-local", base_url=base_url)
    system_text, user_text = _split_system_user(prompt)
    messages = []
    if system_text:
        messages.append({"role": "system", "content": system_text})
    messages.append({"role": "user", "content": user_text})
    kwargs = {
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0,
    }
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    completion = client.chat.completions.create(**kwargs)
    return (completion.choices[0].message.content or "").strip()


_AI_UNICODE_REPLACEMENTS = {
    # Smart quotes
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"',
    "′": "'", "″": '"',
    # Dashes / hyphens
    "‐": "-", "‑": "-", "‒": "-", "–": "-",
    "—": "-", "―": "-",
    # Spaces
    " ": " ", " ": " ", " ": " ", " ": " ", " ": " ",
    # Bullets / ellipsis
    "•": "-", "‣": "-", "◦": "-", "⁃": "-",
    "…": "...",
    # Box-drawing characters used in our prompt template (decorative only)
    "─": "-", "━": "-", "│": "|", "┃": "|",
    "═": "=", "║": "=",
}
_AI_TRANSLATE_TABLE = str.maketrans(_AI_UNICODE_REPLACEMENTS)


def _ascii_strip(s) -> str:
    """For values that must travel in HTTP headers / query params (api_key, model,
    base_url): drop everything that's not ASCII printable. Catches invisible smart
    quotes, zero-width chars, BOMs, and other paste-artifacts that break httpx."""
    if not s:
        return ""
    if not isinstance(s, str):
        s = str(s)
    return "".join(c for c in s if 0x20 <= ord(c) < 0x7F).strip()


def _ai_safe(text: str) -> str:
    """Sanitize text headed for an AI provider's HTTP body.

    Some SDK internals (and HTTP transports for headers) assume ASCII somewhere.
    The body itself is UTF-8, but we've seen 'ascii' codec errors with smart quotes
    on certain Windows + httpx combos. Better to scrub before sending — the AI does
    not need fancy typography.

    Strategy: explicit replacements for common typography, then NFKD-normalize and
    drop any remaining combining marks that don't fold to ASCII. Accented letters
    like é/ñ/ü become their ASCII bases.
    """
    if not isinstance(text, str):
        text = str(text or "")
    text = text.translate(_AI_TRANSLATE_TABLE)
    # Normalize remaining unicode → ASCII where possible
    text = unicodedata.normalize("NFKD", text)
    return text.encode("ascii", "ignore").decode("ascii")


import time as _time


def _is_transient_provider_error(exc: Exception) -> bool:
    """Heuristic: is this error worth retrying?

    Provider 503 / "overloaded" / "high demand" / explicit ServerError → yes.
    Auth, invalid arg, content filtering, daily quota exhausted → no (retrying
    can't fix these).
    """
    err_str = str(exc)
    err_lower = err_str.lower()
    name = type(exc).__name__.lower()
    # Hard quota exhaustion — don't bother retrying, the daily window won't reopen
    # in seconds.
    if "resource_exhausted" in err_lower or "quota exceeded" in err_lower or "exceeded your current quota" in err_lower:
        return False
    if "servererror" in name:
        return True
    if "503" in err_str or "502" in err_str or "504" in err_str:
        return True
    if "unavailable" in err_lower or "overloaded" in err_lower or "high demand" in err_lower:
        return True
    return False


def call_ai(
    provider: str,
    api_key: str,
    prompt: str,
    max_tokens: int = 16000,
    model: str | None = None,
    base_url: str | None = None,
    json_mode: bool = True,
) -> str:
    """Route to the correct AI provider and return raw text.

    Auto-retries up to 3 times on transient provider errors (503 / overloaded /
    high demand) with exponential backoff (1.5s, 4s, 9s). Non-retryable errors
    (auth, invalid request) bubble up immediately.
    """
    provider = (provider or "anthropic").lower()
    # Backward-compat: old frontend sent "claude"
    if provider == "claude":
        provider = "anthropic"
    chosen_model = _ascii_strip(model) or PROVIDER_DEFAULTS.get(provider, "")
    # Defensively scrub everything that touches an HTTP header, query param, or body.
    # Prevents 'ascii' codec errors when users paste API keys / model names that have
    # invisible smart quotes, zero-width chars, or non-breaking spaces.
    api_key = _ascii_strip(api_key)
    base_url = _ascii_strip(base_url) if base_url else None
    prompt = _ai_safe(prompt)

    def _dispatch():
        if provider == "anthropic":
            return _call_anthropic(api_key, chosen_model, prompt, max_tokens, json_mode, base_url)
        if provider == "gemini":
            return _call_gemini(api_key, chosen_model, prompt, max_tokens, json_mode, base_url)
        if provider in OPENAI_COMPATIBLE_BASE_URLS:
            url = base_url or OPENAI_COMPATIBLE_BASE_URLS[provider]
            return _call_openai_compatible(api_key, chosen_model, prompt, max_tokens, json_mode, url)
        if provider == "openai_compatible":
            if not base_url:
                raise ValueError("openai_compatible provider requires a base_url.")
            if not chosen_model:
                raise ValueError("openai_compatible provider requires a model name.")
            return _call_openai_compatible(api_key, chosen_model, prompt, max_tokens, json_mode, base_url)
        raise ValueError(f"Unknown AI provider: {provider}")

    # Retry loop for transient provider errors (Gemini "high demand" / OpenAI 503 / etc.)
    backoffs = (1.5, 4.0, 9.0)  # wait between attempts; total ~14.5s worst case
    last_exc = None
    for attempt, wait in enumerate([0.0] + list(backoffs)):
        if wait > 0:
            print(f"[call_ai] transient error, retrying in {wait}s (attempt {attempt+1}/{len(backoffs)+1})")
            _time.sleep(wait)
        try:
            return _dispatch()
        except Exception as e:
            last_exc = e
            if not _is_transient_provider_error(e):
                raise
            # else: fall through to next backoff
    # Exhausted retries — surface the last error.
    raise last_exc


def _provider_error_response(exc: Exception):
    """Map any provider's auth / rate-limit / generic error to a clean HTTP response."""
    # Always print the full traceback to the server log — silent 500s with one-line
    # messages make these very hard to debug.
    traceback.print_exc()
    err_str = str(exc)
    err_lower = err_str.lower()
    name = type(exc).__name__.lower()
    # Auth-ish
    if "authentication" in name or "API_KEY_INVALID" in err_str or "Incorrect API key" in err_str or "401" in err_str:
        return jsonify({"error": "Invalid API key. Please check the key for the selected provider."}), 401
    # Daily quota exhausted (Gemini free tier, OpenAI billing limits, etc.) — distinct
    # from per-minute rate limiting because the user can't just "wait a moment".
    if "resource_exhausted" in err_lower or "quota exceeded" in err_lower or "exceeded your current quota" in err_lower:
        # Try to surface model + tip
        msg = ("Daily quota exhausted for this model. Options: "
               "(1) switch the Model to a higher-quota one (e.g. 'gemini-2.5-flash-lite' has 1,500/day on Gemini free tier vs. 20/day for 'gemini-2.5-flash'), "
               "(2) switch Provider to Anthropic / OpenAI in the dropdown, "
               "(3) enable Fast mode to use ~6× fewer calls per tailor.")
        return jsonify({"error": msg}), 429
    # Per-minute rate-limit-ish
    if "ratelimit" in name or "rate_limit" in err_lower or "429" in err_str:
        return jsonify({"error": "Rate limited. Please wait a moment and try again."}), 429
    # Provider overload / temporarily unavailable (Gemini's "model experiencing high demand", OpenAI 503, etc.)
    if "unavailable" in err_lower or "503" in err_str or "overloaded" in err_lower or "high demand" in err_lower or "servererror" in name:
        return jsonify({"error": "The AI provider is overloaded right now. Wait 30 seconds and retry, or pick a different model/provider in the dropdown."}), 503
    # Connection-ish (Ollama not running, bad base_url, etc.)
    if "connection" in name or "Connection" in err_str or "ECONNREFUSED" in err_str:
        return jsonify({"error": "Could not reach the AI provider. Check the base URL or that the local server is running."}), 502
    return jsonify({"error": f"Something went wrong: {err_str}"}), 500


def _validate_resume_shape(data: dict) -> None:
    """Sanity-check AI response before downstream code indexes into it."""
    if not isinstance(data, dict):
        raise ValueError("AI did not return a JSON object.")
    if "sections" not in data and "experience" not in data and "summary" not in data:
        # Truncated output is the usual cause — surface a clear hint.
        raise ValueError("AI response is missing the resume body. The output may have been truncated; try again.")


def _section_headings(resume: dict) -> list[str]:
    """Lower-cased ordered list of section headings present in a resume JSON."""
    if not isinstance(resume, dict):
        return []
    return [
        (s.get("heading", "") or "").strip().lower()
        for s in (resume.get("sections") or [])
        if isinstance(s, dict)
    ]


def _has_all_sections(after: dict, before: dict) -> bool:
    """True if `after` retains every section heading present in `before`."""
    before_set = set(h for h in _section_headings(before) if h)
    after_set = set(h for h in _section_headings(after) if h)
    return before_set.issubset(after_set)


def _section_count(resume: dict) -> int:
    """Total non-empty sections in a resume."""
    secs = resume.get("sections") if isinstance(resume, dict) else None
    return len(secs) if isinstance(secs, list) else 0


def _do_tailor_call(api_key, resume_text, prompt_text, jd_text, provider,
                    detected_sections, model, base_url, extra_instruction=""):
    prompt = _build_tailor_user_msg(resume_text, prompt_text, jd_text, detected_sections)
    if extra_instruction:
        prompt += "\n\n" + extra_instruction
    raw = call_ai(provider, api_key, prompt, max_tokens=16000, model=model, base_url=base_url, json_mode=True)
    raw = _strip_code_fences(raw)
    try:
        data = _safe_json_loads(raw)
    except json.JSONDecodeError as e:
        try:
            # Vercel's function dir is read-only; only /tmp is writable.
            base = "/tmp" if os.environ.get("VERCEL") else os.path.dirname(__file__)
            dump_path = os.path.join(base, "last_failed_ai_response.txt")
            with open(dump_path, "w", encoding="utf-8") as f:
                f.write(raw)
            print(f"[debug] dumped failing AI response to {dump_path}")
        except Exception:
            pass
        raise ValueError(f"AI returned invalid JSON even after repair. First 500 chars: {raw[:500]!r}") from e
    _validate_resume_shape(data)
    return data


def tailor_resume(
    api_key: str,
    resume_text: str,
    prompt_text: str,
    jd_text: str,
    provider: str = "anthropic",
    detected_sections: list[str] | None = None,
    model: str | None = None,
    base_url: str | None = None,
) -> dict:
    """Call AI to tailor the resume. Returns structured JSON.

    If the model drops sections that were detected in the original resume, retries
    once with an explicit instruction listing each missing section by name.
    """
    data = _do_tailor_call(api_key, resume_text, prompt_text, jd_text, provider,
                           detected_sections, model, base_url)

    # If we know what sections the original had, verify the AI kept them.
    if detected_sections:
        expected = {s.strip().lower() for s in detected_sections if s.strip()}
        got = set(_section_headings(data))
        missing = sorted(expected - got)
        if missing:
            print(f"[tailor-guard] AI dropped sections {missing}; retrying with explicit list.")
            extra = (
                "MANDATORY FIX — the previous attempt dropped these sections. You MUST include "
                "every one of them in the JSON you return now, in the SAME order they appear in "
                f"the original resume: {', '.join(detected_sections)}.\n"
                f"You especially must include: {', '.join(missing)}.\n"
                "Do NOT merge them, do NOT skip any. If a section had no JD-relevant content, "
                "still include it with whatever the original resume contained."
            )
            try:
                retry = _do_tailor_call(api_key, resume_text, prompt_text, jd_text, provider,
                                        detected_sections, model, base_url, extra_instruction=extra)
                # Use the retry only if it has at least as many sections AND restores the missing ones.
                retry_got = set(_section_headings(retry))
                if (expected - retry_got) == set() or len(retry_got) > len(got):
                    data = retry
                else:
                    print(f"[tailor-guard] retry still missing sections; keeping first response.")
            except Exception as e:
                print(f"[tailor-guard] retry failed ({e}); keeping first response.")

    return data


# ─────────────────────────────────────────────
# ATS scoring + auto-improvement
# ─────────────────────────────────────────────

ATS_TARGET_SCORE = 85
ATS_FLOOR_SCORE = 80           # anything below = aggressive mistake-diagnosis mode
ATS_MAX_IMPROVE_PASSES = 4     # cap on rewrite attempts
ATS_NO_PROGRESS_STOP = 2       # bail out if N consecutive passes don't raise the score


def _build_ats_score_prompt(resume_json: dict, jd_text: str) -> str:
    """Prompt that asks the model to ACT AS AN ATS and score keyword coverage."""
    system = """You are an Applicant Tracking System (ATS) keyword scorer.
You will be given a tailored resume (as JSON) and a job description.

Your job:
1. Extract every distinct technical noun, framework, tool, methodology, certification,
   domain term, and required soft skill from the job description. These are the
   "JD keywords" the ATS will scan for. Treat multi-word terms (e.g. "event-driven
   architecture", "core web vitals") as single keywords. Ignore generic filler
   words ("team", "fast-paced", "collaborate" alone, etc.). Aim for 25-50 keywords.

2. For each JD keyword, check whether it appears (case-insensitive, hyphen/space-
   insensitive, obvious surface variants) anywhere in the resume JSON: name, title,
   contact, section headings, content, items, bullets — every text field.

3. Compute coverage = matched_count / total_count, expressed as an integer 0-100.

4. Identify missing keywords that would most likely raise the score if surfaced.

Return ONLY a JSON object in this exact shape, no prose, no fences:

{
  "score": <integer 0-100>,
  "matched_keywords": ["..."],
  "missing_keywords": ["..."],
  "notes": "one short sentence on the biggest gap, or empty string"
}"""
    user = f"""=== TAILORED RESUME (JSON) ===
{json.dumps(resume_json, ensure_ascii=False)}

=== JOB DESCRIPTION ===
{jd_text}

Score this resume now. Return ONLY the JSON object."""
    return f"{system}\n\n{user}"


def score_resume_ats(
    provider: str,
    api_key: str,
    resume_json: dict,
    jd_text: str,
    model: str | None = None,
    base_url: str | None = None,
) -> dict:
    """Run a single ATS-scoring pass. Returns {score, matched_keywords, missing_keywords, notes}."""
    prompt = _build_ats_score_prompt(resume_json, jd_text)
    raw = call_ai(provider, api_key, prompt, max_tokens=2000, model=model, base_url=base_url, json_mode=True)
    raw = _strip_code_fences(raw)
    parsed = _safe_json_loads(raw)
    if not isinstance(parsed, dict) or "score" not in parsed:
        raise ValueError("ATS scorer returned malformed JSON.")
    # Normalize
    parsed["score"] = int(parsed.get("score", 0))
    parsed["matched_keywords"] = list(parsed.get("matched_keywords") or [])
    parsed["missing_keywords"] = list(parsed.get("missing_keywords") or [])
    parsed["notes"] = str(parsed.get("notes") or "")
    return parsed


def _build_ats_improve_prompt(
    resume_json: dict,
    jd_text: str,
    ats_report: dict,
    original_resume_text: str,
    target: int,
    aggressive: bool,
) -> str:
    """Prompt that rewrites the tailored resume to push the ATS score above the target.

    aggressive=True triggers a "diagnose mistakes first" mode used when the score is
    below the floor (default 80). The model is given a checklist of common ATS failure
    modes and instructed to apply ALL applicable fixes in a single revision.
    """
    missing = ", ".join(ats_report.get("missing_keywords") or []) or "(none listed)"
    matched = ", ".join((ats_report.get("matched_keywords") or [])[:30]) or "(none yet)"

    base_constraints = f"""You are revising a tailored resume to raise its ATS keyword-coverage score to {target}% or higher.

CONSTRAINTS — non-negotiable:
- NEVER fabricate experience, technologies, projects, metrics, dates, or companies.
- Only surface keywords the candidate genuinely has, evidenced by the ORIGINAL resume.
- If a missing keyword is not supported by the original resume, leave it out.
- Keep the same JSON structure (same section types, same heading order).
- Keep the candidate's name, title, and contact unchanged.
- Do not invent new sections; only rewrite existing fields."""

    standard_guidance = """How to raise the score legitimately:
- Reword existing bullets to use the JD's exact terminology (e.g., "K8s" -> "Kubernetes"
  if Kubernetes is what JD uses; "C Sharp" -> "C#"; "message-based" -> "event-driven").
- Move keywords from buried sentences into Skills (as new entries in existing categories)
  and the Summary's first 25 words.
- Expand acronyms once: "CI/CD (Continuous Integration / Continuous Deployment)".
- Rename existing Skills categories to mirror JD section labels exactly."""

    aggressive_diagnosis = f"""THE CURRENT SCORE IS BELOW THE FLOOR ({ATS_FLOOR_SCORE}%). THIS IS A REMEDIATION PASS.

Before rewriting, silently diagnose ALL of the following common ATS failure modes,
then apply EVERY fix that applies. Do not output the diagnosis — only the corrected JSON.

DIAGNOSTIC CHECKLIST (run through every item):
1. TERMINOLOGY MISMATCH — Resume uses synonyms / older names where JD uses a specific term.
   Fix: replace with the JD's exact term wherever the candidate's experience supports it.
   Examples: "ML" vs "Machine Learning", "K8s" vs "Kubernetes", "AWS Lambda" vs "Serverless",
   "REST APIs" vs "RESTful Services", "CI/CD" vs "Continuous Integration / Continuous Deployment".

2. BURIED KEYWORDS — A skill appears once in a buried bullet but not in Skills section.
   Fix: surface it in Skills (under the most relevant existing category) AND keep the bullet.
   Required keywords should appear at least 2x: once in Skills, once in Experience or Summary.

3. SUMMARY KEYWORD GAP — JD's top 2-3 keywords are missing from Summary's first 25 words.
   Fix: rewrite Summary opening to front-load those keywords (without fabricating).

4. SKILLS CATEGORY LABEL DRIFT — Skills categories use generic names ("Backend", "Cloud")
   when JD uses richer phrasing ("Distributed Systems & Microservices", "AWS Cloud Architecture").
   Fix: rename categories to mirror JD section labels exactly.

5. ACRONYM-FORM MISMATCH — JD uses one form (acronym OR expanded), resume uses the other.
   Fix: include BOTH forms once: "Continuous Integration / Continuous Deployment (CI/CD)".

6. HYPHEN / SPACE MISMATCH — "event driven" vs "event-driven", "fullstack" vs "full-stack".
   Fix: match the JD's exact spacing/hyphenation.

7. PRIORITY-STACK INVERSION — JD's #1 priority is buried in the resume's last role.
   Fix: surface it earlier — Summary, Skills, and the most recent role's first bullet.

8. SOFT-SKILL OMISSION — JD names mentoring / cross-functional / ownership; resume implies
   them but never uses the words. Fix: use the explicit JD verbs in bullets.

9. MISSING DOMAIN TERMS — JD mentions a regulated industry / scale signal / framework
   (HIPAA, GDPR, "5M+ users", "real-time"); resume has the experience but doesn't say so.
   Fix: surface the domain term where the candidate actually worked there.

10. WEAK-VERB BULLETS — Bullets start with "Helped", "Worked on", "Was responsible for".
    Fix: replace with strong verbs that match JD tone (Architected / Owned / Drove /
    Optimized / Productionized / Mentored).

PRIORITY ORDER for this pass: (1) terminology fixes, (2) buried keywords surfaced to Skills,
(3) Summary front-loading, (4) Skills category renaming. These four typically move scores
from the 60s/70s into the 85+ range without touching the candidate's actual experience.

After applying fixes, mentally re-score against the JD keyword list. If the projected
score is still below {target}%, apply more fixes from the list until it isn't."""

    system = base_constraints + "\n\n" + (aggressive_diagnosis if aggressive else standard_guidance)
    system += "\n\nReturn ONLY the revised resume JSON in the same schema, no prose, no fences."

    user = f"""=== ORIGINAL RESUME (source of truth — do not exceed what's here) ===
{original_resume_text}

=== CURRENT TAILORED RESUME (JSON) ===
{json.dumps(resume_json, ensure_ascii=False)}

=== JOB DESCRIPTION ===
{jd_text}

=== ATS REPORT ===
Current score: {ats_report.get('score')}  (target {target}%, floor {ATS_FLOOR_SCORE}%)
Already matched (do not lose these): {matched}
Missing keywords (surface these where the original resume genuinely supports them): {missing}
Notes: {ats_report.get('notes', '')}

Rewrite the tailored resume JSON now. Return ONLY the revised JSON object."""
    return f"{system}\n\n{user}"


def improve_resume_for_ats(
    provider: str,
    api_key: str,
    resume_json: dict,
    jd_text: str,
    ats_report: dict,
    original_resume_text: str,
    target: int = ATS_TARGET_SCORE,
    model: str | None = None,
    base_url: str | None = None,
    aggressive: bool = False,
) -> dict:
    """Single rewrite pass to incorporate missing keywords without fabrication.

    aggressive=True activates the diagnose-then-fix prompt used when score < ATS_FLOOR_SCORE.
    """
    prompt = _build_ats_improve_prompt(
        resume_json, jd_text, ats_report, original_resume_text, target, aggressive,
    )
    raw = call_ai(provider, api_key, prompt, max_tokens=16000, model=model, base_url=base_url, json_mode=True)
    raw = _strip_code_fences(raw)
    data = _safe_json_loads(raw)
    _validate_resume_shape(data)
    return data


def score_and_improve_resume_ats(
    provider: str,
    api_key: str,
    resume_json: dict,
    jd_text: str,
    original_resume_text: str,
    target: int = ATS_TARGET_SCORE,
    aggressive: bool = False,
    model: str | None = None,
    base_url: str | None = None,
) -> tuple[dict, dict]:
    """One-shot: score AND (if low) return an improved resume in a single AI call.

    Saves a full round-trip vs. calling score_resume_ats then improve_resume_for_ats.
    Returns (improved_or_unchanged_resume_json, ats_report).
    """
    missing_hint = ", ".join((resume_json.get("missing_keywords") or [])[:0])  # not used; kept for symmetry
    score_section = _build_ats_score_prompt(resume_json, jd_text)
    improve_section = _build_ats_improve_prompt(
        resume_json, jd_text,
        {"score": "(unknown — you compute it)", "matched_keywords": [], "missing_keywords": [], "notes": ""},
        original_resume_text, target, aggressive,
    )

    combined_system = f"""You will perform TWO tasks in one response and return a single JSON object.

TASK 1 — SCORE THE RESUME:
{score_section.split('Return ONLY')[0]}

TASK 2 — IF SCORE < {target}, RETURN AN IMPROVED RESUME:
{improve_section.split('Return ONLY')[0]}

OUTPUT — return ONLY this JSON object, no prose, no fences:

{{
  "score": <integer 0-100>,
  "matched_keywords": ["..."],
  "missing_keywords": ["..."],
  "notes": "one short sentence on the biggest gap, or empty string",
  "improved_resume": <revised resume JSON if score < {target}, else null>
}}

If the score is already >= {target}, set improved_resume to null and skip the rewrite.
If you do rewrite, the improved_resume MUST follow the same schema as the input
resume (same section types, same heading order, name/title/contact unchanged)."""
    user = f"""=== ORIGINAL RESUME (source of truth) ===
{original_resume_text}

=== CURRENT TAILORED RESUME (JSON) ===
{json.dumps(resume_json, ensure_ascii=False)}

=== JOB DESCRIPTION ===
{jd_text}

Return ONLY the combined JSON object."""
    prompt = f"{combined_system}\n\n{user}"
    raw = call_ai(provider, api_key, prompt, max_tokens=16000, model=model, base_url=base_url, json_mode=True)
    raw = _strip_code_fences(raw)
    parsed = _safe_json_loads(raw)
    if not isinstance(parsed, dict) or "score" not in parsed:
        raise ValueError("Combined ATS scorer returned malformed JSON.")

    ats = {
        "score": int(parsed.get("score", 0)),
        "matched_keywords": list(parsed.get("matched_keywords") or []),
        "missing_keywords": list(parsed.get("missing_keywords") or []),
        "notes": str(parsed.get("notes") or ""),
    }
    improved = parsed.get("improved_resume")
    if improved and isinstance(improved, dict):
        try:
            _validate_resume_shape(improved)
            # Reject any improvement that drops sections we previously had.
            if not _has_all_sections(improved, resume_json):
                print(f"[improve-guard] rejecting rewrite: lost sections "
                      f"({_section_count(resume_json)} -> {_section_count(improved)}). "
                      f"keeping previous version.")
                return resume_json, ats
            # Reject if a section that previously had items is now empty.
            before_by_h = {(s.get("heading","") or "").strip().lower(): s
                           for s in (resume_json.get("sections") or []) if isinstance(s, dict)}
            after_by_h = {(s.get("heading","") or "").strip().lower(): s
                          for s in (improved.get("sections") or []) if isinstance(s, dict)}
            for h, before_sec in before_by_h.items():
                after_sec = after_by_h.get(h)
                if not after_sec:
                    continue
                # Items / content presence check
                bef_items = before_sec.get("items") or before_sec.get("content")
                aft_items = after_sec.get("items") or after_sec.get("content")
                if bef_items and not aft_items:
                    print(f"[improve-guard] rejecting rewrite: section {h!r} "
                          f"became empty. keeping previous version.")
                    return resume_json, ats
            return improved, ats
        except Exception as e:
            print(f"[improve-guard] rejecting rewrite due to validation error: {e}")
    return resume_json, ats


def tailor_with_ats_target(
    api_key: str,
    resume_text: str,
    prompt_text: str,
    jd_text: str,
    provider: str,
    detected_sections: list[str] | None,
    model: str | None,
    base_url: str | None,
    target: int = ATS_TARGET_SCORE,
    fast_mode: bool = False,
) -> tuple[dict, dict | None]:
    """Tailor → score → (if low) improve → re-score. Returns (final_resume_json, ats_report_or_None).

    fast_mode=True: skip the score+improve loop entirely and return the first tailor pass.
    Saves 1-N AI round-trips at the cost of an unverified ATS score.

    The ATS pass is best-effort: if it fails for any reason, we still return the tailored
    resume — ATS scoring should never block delivery of the resume itself.
    """
    result = tailor_resume(
        api_key, resume_text, prompt_text, jd_text, provider, detected_sections,
        model=model, base_url=base_url,
    )

    if fast_mode:
        # No ATS pass — fastest path. Score is unknown.
        return result, {"score": None, "fast_mode": True, "target": target, "floor": ATS_FLOOR_SCORE,
                        "passes": 0, "history": [], "matched_keywords": [], "missing_keywords": [],
                        "notes": "Fast mode — ATS scoring skipped."}

    # Snapshot of the initial tailored resume (pre-remediation) so the frontend can diff.
    initial_result = json.loads(json.dumps(result))

    ats: dict | None = None
    try:
        ats = score_resume_ats(provider, api_key, result, jd_text, model=model, base_url=base_url)
        history = [ats["score"]]
        initial_ats = json.loads(json.dumps(ats))
        # Keep the best result we've seen — if a remediation pass somehow makes things
        # worse, we won't ship the regression.
        best_result, best_ats = result, ats
        no_progress = 0
        passes = 0

        while best_ats["score"] < target and passes < ATS_MAX_IMPROVE_PASSES:
            aggressive = best_ats["score"] < ATS_FLOOR_SCORE
            # COMBINED score+improve in ONE call (saves a round-trip per pass).
            improved, new_ats = score_and_improve_resume_ats(
                provider, api_key, best_result, jd_text, resume_text, target, aggressive,
                model=model, base_url=base_url,
            )
            history.append(new_ats["score"])
            passes += 1

            if new_ats["score"] > best_ats["score"]:
                best_result, best_ats = improved, new_ats
                no_progress = 0
            else:
                no_progress += 1
                if no_progress >= ATS_NO_PROGRESS_STOP:
                    # Two passes without improvement — we've plateaued. Stop burning tokens.
                    break

        result, ats = best_result, best_ats
        ats["target"] = target
        ats["floor"] = ATS_FLOOR_SCORE
        ats["passes"] = passes
        ats["history"] = history
        ats["initial_score"] = initial_ats.get("score")
        # Only attach the snapshot if the remediation actually ran AND the result changed.
        if passes > 0 and result is not initial_result:
            ats["initial_resume"] = initial_result
    except Exception as e:
        # Don't fail the whole request just because the scorer choked.
        ats = {"score": None, "error": str(e), "target": target, "floor": ATS_FLOOR_SCORE,
               "passes": 0, "history": [], "matched_keywords": [], "missing_keywords": [], "notes": ""}

    return result, ats


# ─────────────────────────────────────────────
# JD keyword preview (pre-submit)
# ─────────────────────────────────────────────

# Filler words / generic terms that should never count as ATS keywords.
_JD_KW_STOPWORDS = {
    "the", "and", "or", "a", "an", "of", "in", "on", "for", "to", "with", "as", "by",
    "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do",
    "does", "did", "will", "would", "should", "could", "can", "may", "might", "must",
    "this", "that", "these", "those", "there", "their", "they", "them", "we", "you",
    "your", "our", "us", "i", "he", "she", "it", "his", "her", "its",
    "team", "teams", "role", "roles", "company", "companies", "candidate", "candidates",
    "experience", "experienced", "knowledge", "ability", "abilities", "skills", "skill",
    "work", "working", "works", "worked", "year", "years", "month", "months", "week",
    "day", "days", "time", "times", "good", "great", "excellent", "strong", "solid",
    "passionate", "fast-paced", "fast", "paced", "dynamic", "highly", "very", "well",
    "across", "within", "throughout", "while", "during", "including", "such",
    "responsibilities", "responsibility", "qualifications", "qualified", "preferred",
    "required", "requires", "requirement", "requirements", "must", "plus", "bonus",
    "opportunity", "position", "job", "jobs", "career", "careers", "applicant",
    "minimum", "maximum", "etc", "ie", "eg", "e.g.", "i.e.", "and/or",
    "english", "language", "communication", "communicate", "communicating",
    "collaborate", "collaboration", "collaborative",
    "develop", "developing", "development", "developer", "developers",
    "build", "building", "builder", "builders", "make", "making", "create", "creating",
    "ensure", "ensuring", "support", "supporting", "help", "helping", "use", "using",
    "based", "level", "senior", "junior", "mid", "lead", "principal", "staff",
    "month", "months", "week", "weeks", "day", "days", "hour", "hours", "etc.",
    "etc", "us", "usa", "united", "states", "remote", "hybrid", "onsite", "on-site",
    "fulltime", "full-time", "parttime", "part-time", "contract", "contractor",
}

# Strong "this is a real keyword" hints: tokens that contain these substrings (case-insensitive)
# get an automatic boost so they survive the cut even if they appear only once.
_JD_KW_BOOSTERS = (
    "++", "#", ".js", ".net", "sql", "ai", "ml", "api", "sdk", "ci/cd", "k8s",
    "aws", "gcp", "azure", "saas", "ios", "ux", "ui",
)


def extract_jd_keywords_heuristic(jd_text: str, top_n: int = 40) -> list[dict]:
    """Cheap, no-AI keyword extraction. Used for free preview before user spends tokens.

    Returns a list of {keyword, count} dicts, ordered by frequency. Picks up:
    - Capitalized multi-word terms (likely tech/framework names)
    - Hyphenated compound terms
    - Acronyms (2-5 uppercase letters)
    - Tokens containing booster substrings (.js, ++, #, etc.)
    """
    text = jd_text or ""
    if not text.strip():
        return []

    counts: dict[str, int] = {}

    def _bump(token: str, weight: int = 1):
        key = token.strip().strip(".,;:()[]{}\"'").strip()
        if not key:
            return
        if key.lower() in _JD_KW_STOPWORDS:
            return
        if len(key) < 2:
            return
        # Don't keep pure numbers
        if key.replace(",", "").replace(".", "").isdigit():
            return
        counts[key] = counts.get(key, 0) + weight

    # 1) Multi-word capitalized phrases (e.g., "Continuous Integration", "Core Web Vitals")
    for m in re.finditer(r"\b(?:[A-Z][a-zA-Z0-9+#./-]+)(?:\s+[A-Z][a-zA-Z0-9+#./-]+){1,3}\b", text):
        _bump(m.group(0), weight=3)

    # 2) Acronyms (2-5 uppercase letters / digits)
    for m in re.finditer(r"\b[A-Z][A-Z0-9]{1,4}\b", text):
        _bump(m.group(0), weight=2)

    # 3) Hyphenated tech words (event-driven, micro-services, full-stack)
    for m in re.finditer(r"\b[a-zA-Z][a-zA-Z0-9]+(?:-[a-zA-Z][a-zA-Z0-9]+){1,3}\b", text):
        _bump(m.group(0), weight=2)

    # 4) Tokens containing strong booster substrings (.js, ++, #, etc.)
    for m in re.finditer(r"[A-Za-z][A-Za-z0-9+#./-]*", text):
        tok = m.group(0)
        low = tok.lower()
        if any(b in low for b in _JD_KW_BOOSTERS):
            _bump(tok, weight=2)

    # 5) Single capitalized words that aren't sentence-starts: harder to do robustly
    #    without sentence segmentation. Keep light: tokens that begin with capital and
    #    contain a digit or special char are likely product names (e.g., "Postgres", "MongoDB").
    for m in re.finditer(r"\b[A-Z][a-zA-Z]*[A-Z0-9][a-zA-Z0-9]*\b", text):
        _bump(m.group(0), weight=1)

    # Sort by count desc, then keyword asc; cap at top_n.
    items = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0].lower()))
    return [{"keyword": k, "count": c} for k, c in items[:top_n]]


def extract_jd_keywords_ai(
    provider: str,
    api_key: str,
    jd_text: str,
    model: str | None = None,
    base_url: str | None = None,
) -> list[dict]:
    """Ask the AI to extract the same keywords the ATS scorer would look for.
    Uses the same definition as score_resume_ats so the preview matches reality."""
    system = """You are an Applicant Tracking System (ATS) keyword extractor.
From a job description, extract the 25-50 distinct technical/domain keywords an ATS
would scan for: technologies, frameworks, tools, methodologies, certifications,
required soft skills (like "mentoring"), and domain terms. Treat multi-word terms
("event-driven architecture", "core web vitals") as single keywords. Ignore filler
words ("team", "fast-paced", "passionate", generic verbs).

Return ONLY a JSON object with this shape, no prose, no fences:
{
  "keywords": [
    {"keyword": "TypeScript",  "priority": "required",  "category": "language"},
    {"keyword": "Kubernetes",  "priority": "preferred", "category": "devops"}
  ]
}

priority is one of: "required", "preferred", "nice_to_have"
category is one of: "language", "framework", "tool", "platform", "methodology",
                    "domain", "soft_skill", "certification", "data", "other"
"""
    user = f"=== JOB DESCRIPTION ===\n{jd_text}\n\nReturn ONLY the JSON object."
    raw = call_ai(provider, api_key, f"{system}\n\n{user}",
                  max_tokens=1500, model=model, base_url=base_url, json_mode=True)
    raw = _strip_code_fences(raw)
    data = _safe_json_loads(raw)
    kws = data.get("keywords") if isinstance(data, dict) else data
    if not isinstance(kws, list):
        raise ValueError("AI returned malformed keyword list.")
    cleaned = []
    seen = set()
    for k in kws:
        if isinstance(k, str):
            kw = k.strip()
            if kw and kw.lower() not in seen:
                seen.add(kw.lower())
                cleaned.append({"keyword": kw, "priority": "required", "category": "other"})
        elif isinstance(k, dict):
            kw = (k.get("keyword") or "").strip()
            if kw and kw.lower() not in seen:
                seen.add(kw.lower())
                cleaned.append({
                    "keyword": kw,
                    "priority": (k.get("priority") or "required").lower(),
                    "category": (k.get("category") or "other").lower(),
                })
    return cleaned


# ─────────────────────────────────────────────
# PDF generation
# ─────────────────────────────────────────────

_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _clean_text(text) -> str:
    """Replace Unicode characters unsupported by Helvetica with ASCII equivalents,
    and unwrap markdown-style links so ATS sees plain text + URL.

    Final fallback: anything that's still beyond Latin-1 (which is what fpdf2's
    built-in Helvetica supports) gets stripped, so generation never crashes with
    'ascii'/'latin-1' codec errors."""
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    # Unwrap [display](url) -> "display (url)" or just the URL if display == url.
    def _unwrap(m):
        display, url = m.group(1).strip(), m.group(2).strip()
        if display == url or display.lower() == url.lower():
            return url
        return f"{display} ({url})"
    text = _MD_LINK_RE.sub(_unwrap, text)
    cleaned = (text
        .replace("\u2013", "-")   # en dash
        .replace("\u2014", "-")   # em dash
        .replace("\u2018", "'")   # left single quote
        .replace("\u2019", "'")   # right single quote
        .replace("\u201c", '"')   # left double quote
        .replace("\u201d", '"')   # right double quote
        .replace("\u2022", "-")   # bullet
        .replace("\u2026", "...")  # ellipsis
        .replace("\u00a0", " ")   # non-breaking space
        .replace("\u2010", "-")   # hyphen
        .replace("\u2011", "-")   # non-breaking hyphen
        .replace("\u2012", "-")   # figure dash
        .replace("\u2015", "-")   # horizontal bar
        .replace("\u2027", "-")   # hyphenation point
        .replace("\u2032", "'")   # prime
        .replace("\u2033", '"')   # double prime
    )
    # Final fallback: drop anything still beyond Latin-1 so fpdf2's Helvetica
    # never crashes with a codec error. Accented Latin-1 chars (\u00e9, \u00f1, \u00fc, ...)
    # survive; truly exotic glyphs become "?".
    return cleaned.encode("latin-1", "replace").decode("latin-1")


# ─────────────────────────────────────────────
# PDF Templates
# ─────────────────────────────────────────────

PDF_TEMPLATES = {
    "classic_blue": {
        "accent": (37, 99, 165),
        "dark": (26, 26, 26),
        "body": (51, 51, 51),
        "gray": (102, 102, 102),
        "name_align": "C",
        "name_size": 20,
        "header_transform": "upper",
        "header_size": 12,
        "header_line_color": (37, 99, 165),
        "skill_bullet": "-",
        "contact_align": "C",
    },
    "modern_green": {
        "accent": (107, 142, 35),
        "dark": (26, 26, 26),
        "body": (51, 51, 51),
        "gray": (102, 102, 102),
        "name_align": "L",
        "name_size": 24,
        "header_transform": "title",
        "header_size": 14,
        "header_line_color": (200, 200, 200),
        "skill_bullet": "-",
        "contact_align": "L",
    },
    "universal": {
        "accent": (50, 50, 50),
        "dark": (30, 30, 30),
        "body": (45, 45, 45),
        "gray": (110, 110, 110),
        "name_align": "L",
        "name_size": 22,
        "header_transform": "upper",
        "header_size": 11,
        "header_line_color": (180, 180, 180),
        "skill_bullet": "-",
        "contact_align": "L",
    },
    # Mirrors the layout used by Sudderick Matthew's reference resume PDF:
    # large regular-weight name, stacked contact lines, thin grey rule ABOVE each
    # section heading, regular-weight title-case headings, regular-weight job titles,
    # bullet glyphs with **bold** inline keyword highlighting.
    "minimal_clean": {
        "accent":             (35, 35, 35),
        "dark":               (25, 25, 25),
        "body":               (50, 50, 50),
        "gray":               (110, 110, 110),
        "name_align":         "L",
        "name_size":          26,
        "name_bold":          False,           # NEW: name rendered in regular weight
        "contact_align":      "L",
        "contact_stacked":    True,            # NEW: each contact part on its own line
        "header_transform":   "title",
        "header_size":        13,
        "header_bold":        False,           # NEW: heading in regular weight
        "header_line_color":  (210, 210, 210),
        "header_line_position":"above",        # NEW: rule sits ABOVE the heading text
        "skill_bullet":       "-",
        "job_title_bold":     False,           # NEW: job title regular weight
        "job_title_size":     13,              # NEW: larger than other templates
        "job_company_separator": " - ",          # bullet glyph between company and location
        "bullet_inline_bold": True,            # NEW: parse **bold** markdown inside bullets
        "bullet_glyph_size":  0.85,
        "simple_list_as_paragraph": True,      # NEW: skills/etc. render as plain paragraph, no bullets
        "education_layout":   "minimal",       # NEW: degree first, school+location grey, dates grey
    },
}


class ResumePDF(FPDF):

    def __init__(self, template="modern_green"):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        cfg = PDF_TEMPLATES.get(template, PDF_TEMPLATES["modern_green"])
        self.cfg = cfg
        self.ACCENT = cfg["accent"]
        self.DARK = cfg["dark"]
        self.BODY = cfg["body"]
        self.GRAY = cfg["gray"]

    def header(self):
        pass  # no header

    # Belt-and-suspenders: every text passing into fpdf2's core renderers gets
    # cleaned, so it doesn't matter if a call site forgot to wrap with _clean_text.
    def cell(self, w=None, h=None, text="", *args, **kwargs):
        if "txt" in kwargs:  # fpdf2 still accepts the deprecated 'txt' kwarg
            kwargs["txt"] = _clean_text(kwargs["txt"])
            return super().cell(w, h, *args, **kwargs)
        return super().cell(w, h, _clean_text(text), *args, **kwargs)

    def multi_cell(self, w=0, h=None, text="", *args, **kwargs):
        if "txt" in kwargs:
            kwargs["txt"] = _clean_text(kwargs["txt"])
            return super().multi_cell(w, h, *args, **kwargs)
        return super().multi_cell(w, h, _clean_text(text), *args, **kwargs)

    def section_header(self, text):
        transform = self.cfg.get("header_transform", "upper")
        display = _clean_text(text).upper() if transform == "upper" else _clean_text(text)
        weight = "B" if self.cfg.get("header_bold", True) else ""
        line_pos = self.cfg.get("header_line_position", "below")  # "above" / "below" / "none"
        self.ln(2)
        if line_pos == "above":
            # Draw the rule first, then a small gap, then the heading text.
            self.set_draw_color(*self.cfg.get("header_line_color", (200, 200, 200)))
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(3)
        self.set_font("Helvetica", weight, self.cfg.get("header_size", 14))
        self.set_text_color(*self.ACCENT)
        self.cell(0, 9, display, new_x="LMARGIN", new_y="NEXT")
        if line_pos == "below":
            self.set_draw_color(*self.cfg.get("header_line_color", (200, 200, 200)))
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def body_text(self, text, size=10):
        self.set_font("Helvetica", "", size)
        self.set_text_color(*self.BODY)
        self.multi_cell(0, 5, _clean_text(text))

    def bullet(self, text):
        # If template enables inline **bold** markdown, route to the rich-text bullet.
        if self.cfg.get("bullet_inline_bold"):
            return self._bullet_inline_bold(text)
        text = _clean_text(text)
        size = 9.5
        line_h = 4.5
        indent = 8
        bullet_x = self.l_margin + indent
        text_x = bullet_x + 5
        text_w = self.w - self.r_margin - text_x

        self.set_font("Helvetica", "", size)
        self.set_text_color(*self.BODY)

        # Manual word-wrap so wrapped lines stay under text_x (hanging indent),
        # rather than snapping back to the left margin like multi_cell does.
        tokens = [t for t in re.split(r"(\s+)", text) if t]
        lines = []
        cur, cur_w = [], 0.0
        for tok in tokens:
            tw = self.get_string_width(tok)
            if cur and (cur_w + tw) > text_w:
                line_text = "".join(cur).rstrip()
                if line_text:
                    lines.append(line_text)
                if tok.isspace():
                    cur, cur_w = [], 0.0
                else:
                    cur, cur_w = [tok], tw
            else:
                cur.append(tok)
                cur_w += tw
        if cur:
            line_text = "".join(cur).rstrip()
            if line_text:
                lines.append(line_text)
        if not lines:
            lines = [""]

        # If the first line wouldn't fit on the current page, force a page
        # break BEFORE drawing the glyph. Otherwise the circle is rendered
        # at the bottom of page N while the cell auto-breaks and the text
        # lands on page N+1, leaving an orphaned bullet.
        if self.get_y() + line_h > self.h - self.b_margin:
            self.add_page()

        # Vertically center the bullet glyph with the cap-middle of the first
        # line of text. fpdf2.circle treats (x, y) as the *center* of the
        # circle, and the cap-middle of 9.5pt Helvetica inside a 4.5mm cell
        # sits at roughly y + line_h * 0.55 — calibrated empirically.
        glyph_cy = self.get_y() + line_h * 0.55
        for i, line in enumerate(lines):
            if i == 0:
                self.circle(bullet_x + 1, glyph_cy, 0.8, style="F")
            self.set_x(text_x)
            self.cell(text_w, line_h, line)
            self.ln(line_h)
        self.ln(0.5)

    def _bullet_inline_bold(self, text):
        """Bullet with hanging indent and **bold** markdown rendered inline.

        Used by templates whose cfg sets `bullet_inline_bold: true`. Word-wraps using
        per-segment widths (regular vs. bold) so a long bold span doesn't break layout.
        """
        text = _clean_text(text)
        size = 9.5
        line_h = 4.7
        indent = 6
        bullet_x = self.l_margin + 2
        text_x = self.l_margin + indent
        text_w = self.w - self.r_margin - text_x
        body_color = self.BODY

        # Tokenize: split on **...** so even-indexed parts are regular, odd-indexed are bold.
        parts = text.split("**")
        tokens = []  # list of (weight, word_or_space)
        for idx, part in enumerate(parts):
            weight = "B" if idx % 2 == 1 else ""
            for chunk in re.split(r"(\s+)", part):
                if chunk:
                    tokens.append((weight, chunk))

        # Word-wrap. Each token's width depends on its weight, so we have to switch fonts
        # while measuring.
        lines = []  # list of [(weight, text), ...]
        cur, cur_w = [], 0.0
        for weight, tok in tokens:
            self.set_font("Helvetica", weight, size)
            tw = self.get_string_width(tok)
            if cur and (cur_w + tw) > text_w:
                # Flush — strip trailing whitespace
                while cur and cur[-1][1].isspace():
                    cur_w -= self.get_string_width(cur[-1][1])  # rough, font may differ
                    cur.pop()
                if cur:
                    lines.append(cur)
                # Skip leading whitespace on the new line
                if tok.isspace():
                    cur, cur_w = [], 0.0
                    continue
                cur, cur_w = [(weight, tok)], tw
            else:
                cur.append((weight, tok))
                cur_w += tw
        if cur:
            while cur and cur[-1][1].isspace():
                cur.pop()
            if cur:
                lines.append(cur)

        # Page-break guard — keep the glyph on the same page as the first
        # text line (see bullet() for rationale).
        if self.get_y() + line_h > self.h - self.b_margin:
            self.add_page()

        # Render. Center the glyph vertically with the cap-middle of the first
        # line — same calibration as bullet() above.
        self.set_text_color(*body_color)
        glyph_r = self.cfg.get("bullet_glyph_size", 0.85)
        glyph_cy = self.get_y() + line_h * 0.55
        for i, line in enumerate(lines):
            if i == 0:
                self.circle(bullet_x + 1, glyph_cy, glyph_r, style="F")
            self.set_x(text_x)
            for weight, tok in line:
                self.set_font("Helvetica", weight, size)
                self.cell(self.get_string_width(tok), line_h, tok)
            self.ln(line_h)
        self.ln(0.5)

    def list_bullet(self, text):
        """Render a bullet for simple lists (certs, awards, etc.) using template bullet style."""
        text = _clean_text(text)
        size = 9.5
        line_h = 4.5
        indent = 8
        glyph = self.cfg.get("skill_bullet", "-")
        glyph_w = 5
        x0 = self.get_x()
        bullet_x = x0 + indent
        text_x = bullet_x + glyph_w
        text_w = self.w - self.r_margin - text_x

        self.set_font("Helvetica", "", size)
        self.set_text_color(*self.BODY)

        tokens = [t for t in re.split(r"(\s+)", text) if t]
        lines = []
        cur, cur_w = [], 0.0
        for tok in tokens:
            tw = self.get_string_width(tok)
            if cur and (cur_w + tw) > text_w:
                line_text = "".join(cur).rstrip()
                if line_text:
                    lines.append(line_text)
                if tok.isspace():
                    cur, cur_w = [], 0.0
                else:
                    cur, cur_w = [tok], tw
            else:
                cur.append(tok)
                cur_w += tw
        if cur:
            line_text = "".join(cur).rstrip()
            if line_text:
                lines.append(line_text)
        if not lines:
            lines = [""]

        # Keep the glyph on the same page as the first text line.
        if self.get_y() + line_h > self.h - self.b_margin:
            self.add_page()

        for i, line in enumerate(lines):
            if i == 0:
                self.set_x(bullet_x)
                self.cell(glyph_w, line_h, glyph)
            else:
                self.set_x(text_x)
            self.cell(text_w, line_h, line)
            self.ln(line_h)
        self.ln(0.5)


def _to_list(val):
    """Safely convert a value to a list. Handles None, dict, str, list."""
    if isinstance(val, list):
        return val
    if isinstance(val, dict):
        return [val]
    if isinstance(val, str):
        return [val]
    return []


def _render_section_text(pdf, section):
    """Render a text section (e.g. Summary, Profile, Objective)."""
    pdf.section_header(section.get("heading", "Summary"))
    pdf.body_text(section.get("content", ""))
    pdf.ln(3)


def _wrap_text(pdf, text, max_width):
    """Split text into lines that fit within max_width using current font."""
    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        test = f"{current_line} {word}".strip()
        if pdf.get_string_width(test) <= max_width:
            current_line = test
        else:
            if current_line:
                lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return lines


def _render_section_skills(pdf, section):
    """Render a skills section with arrow bullets and category: items rows."""
    pdf.section_header(section.get("heading", "Skills"))
    arrow_indent = 6   # space for the ">" arrow
    line_h = 4.5

    for skill in _to_list(section.get("items")):
        if isinstance(skill, str):
            pdf.list_bullet(skill)
            continue
        cat = _clean_text(skill.get("category", ""))
        items = _clean_text(skill.get("items", ""))

        # Arrow bullet
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*pdf.BODY)
        pdf.cell(arrow_indent, 5, pdf.cfg.get("skill_bullet", "-"))

        # Category label (bold)
        pdf.set_font("Helvetica", "B", 9.5)
        label = f"{cat}:  "
        label_w = pdf.get_string_width(label)
        pdf.cell(label_w, 5, label)

        # First line: items start right after the bold category label
        first_line_x = pdf.get_x()
        first_line_width = pdf.w - pdf.r_margin - first_line_x

        # Continuation lines: start from left margin + arrow indent
        cont_x = pdf.l_margin + arrow_indent
        cont_width = pdf.w - pdf.r_margin - cont_x

        # Wrap: first line has less space, continuation lines get full width
        pdf.set_font("Helvetica", "", 9.5)
        words = items.split()
        lines = []
        current_line = ""
        on_first = True
        max_w = first_line_width

        for word in words:
            test = f"{current_line} {word}".strip()
            if pdf.get_string_width(test) <= max_w:
                current_line = test
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
                if on_first:
                    on_first = False
                    max_w = cont_width
        if current_line:
            lines.append(current_line)

        for j, line in enumerate(lines):
            if j == 0:
                pdf.cell(first_line_width, line_h, line, new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.set_x(cont_x)
                pdf.cell(cont_width, line_h, line, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
    pdf.ln(1)


def _render_section_experience(pdf, section):
    """Render an experience section with job entries.

    Default templates: "Company – Job Title" bold on line 1, dates grey on line 2.
    Templates with `minimal_clean`-style cfg: large regular Job Title on line 1,
    grey "Company - Location" on line 2, grey dates on line 3.
    """
    pdf.section_header(section.get("heading", "Experience"))
    cfg = pdf.cfg
    minimal_layout = cfg.get("job_title_size") and not cfg.get("job_title_bold", True)
    title_weight = "B" if cfg.get("job_title_bold", True) else ""
    title_size = cfg.get("job_title_size", 10.5)
    sep = cfg.get("job_company_separator", " - ")

    for job in _to_list(section.get("items")):
        if isinstance(job, str):
            pdf.body_text(job)
            continue
        company = _clean_text(job.get("company", ""))
        title = _clean_text(job.get("job_title", ""))
        dates = _clean_text(job.get("dates", ""))
        location = _clean_text(job.get("location", "") or job.get("context", ""))

        if minimal_layout:
            # Line 1: Job Title (regular weight, larger)
            pdf.set_font("Helvetica", title_weight, title_size)
            pdf.set_text_color(*pdf.DARK)
            pdf.cell(0, 7, title or company, new_x="LMARGIN", new_y="NEXT")

            # Line 2: Company - Location (small grey)
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*pdf.GRAY)
            sub_parts = [p for p in (company, location) if p]
            if sub_parts:
                pdf.cell(0, 5, sep.join(sub_parts), new_x="LMARGIN", new_y="NEXT")

            # Line 3: Dates (small grey)
            if dates:
                pdf.cell(0, 5, dates, new_x="LMARGIN", new_y="NEXT")
        else:
            # Original layout
            pdf.set_font("Helvetica", title_weight, title_size)
            pdf.set_text_color(*pdf.DARK)
            company_title = f"{company}{sep}{title}" if company and title else company or title
            pdf.cell(0, 6, company_title, new_x="LMARGIN", new_y="NEXT")
            if dates:
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*pdf.GRAY)
                pdf.cell(0, 5, dates, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

        # Bullets
        for b in _to_list(job.get("bullets")):
            pdf.bullet(str(b))
        pdf.ln(3)


def _render_section_education(pdf, section):
    """Render an education section (supports multiple entries).

    Default layout: School name (bold), degree as sub-item, dates below.
    `education_layout="minimal"`: Degree headline (regular, larger), then
    "School - Location" in grey, then dates in grey — mirrors the experience block.
    """
    pdf.section_header(section.get("heading", "Education"))
    items = _to_list(section.get("items"))
    layout = pdf.cfg.get("education_layout", "default")
    sep = pdf.cfg.get("job_company_separator", " - ")

    for edu in items:
        if isinstance(edu, str):
            pdf.body_text(edu)
            pdf.ln(2)
            continue
        degree = _clean_text(edu.get("degree", ""))
        school = _clean_text(edu.get("school", ""))
        location = _clean_text(edu.get("location", ""))
        dates = _clean_text(edu.get("dates", ""))

        if layout == "minimal":
            # Headline: Degree
            pdf.set_font("Helvetica", "", pdf.cfg.get("job_title_size", 13))
            pdf.set_text_color(*pdf.DARK)
            pdf.cell(0, 7, degree or school, new_x="LMARGIN", new_y="NEXT")
            # Sub-line: School - Location
            sub_parts = [p for p in (school, location) if p]
            if sub_parts and degree:
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*pdf.GRAY)
                pdf.cell(0, 5, sep.join(sub_parts), new_x="LMARGIN", new_y="NEXT")
            # Dates
            if dates:
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*pdf.GRAY)
                pdf.cell(0, 5, dates, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            continue

        # Default layout
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*pdf.DARK)
        pdf.cell(0, 6, school, new_x="LMARGIN", new_y="NEXT")
        if degree:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*pdf.BODY)
            pdf.set_x(pdf.l_margin + 8)
            pdf.cell(5, 5, pdf.cfg.get("skill_bullet", "-"))
            pdf.cell(0, 5, degree, new_x="LMARGIN", new_y="NEXT")
        if dates:
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*pdf.GRAY)
            pdf.set_x(pdf.l_margin + 13)
            pdf.cell(0, 5, dates, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)


def _render_section_projects(pdf, section):
    """Render a projects section."""
    pdf.section_header(section.get("heading", "Projects"))
    for proj in _to_list(section.get("items")):
        if isinstance(proj, str):
            pdf.bullet(proj)
            continue
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*pdf.DARK)
        name = _clean_text(proj.get("name", ""))
        pdf.cell(0, 6, name, new_x="LMARGIN", new_y="NEXT")

        desc = proj.get("description", "")
        tech = proj.get("technologies", "")
        if desc or tech:
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*pdf.ACCENT)
            if tech:
                pdf.cell(0, 5, _clean_text(tech), new_x="LMARGIN", new_y="NEXT")
            if desc:
                pdf.set_text_color(*pdf.BODY)
                pdf.multi_cell(0, 4.5, _clean_text(desc))
            pdf.ln(1)

        for b in _to_list(proj.get("bullets")):
            pdf.bullet(str(b))
        pdf.ln(3)


def _render_section_simple_list(pdf, section):
    """Render a simple list section (certifications, awards, languages, etc.)."""
    pdf.section_header(section.get("heading", ""))

    # Templates can opt to render simple lists as a plain comma-joined paragraph
    # instead of a bulleted list (matches the reference resume's Skills layout).
    if pdf.cfg.get("simple_list_as_paragraph"):
        items_text = []
        for item in _to_list(section.get("items")):
            if isinstance(item, str):
                items_text.append(item.strip())
            elif isinstance(item, dict):
                items_text.append((item.get("name", "") or item.get("title", "") or "").strip())
        joined = ", ".join(p for p in items_text if p)
        if joined:
            pdf.body_text(joined)
        pdf.ln(2)
        return

    for item in _to_list(section.get("items")):
        if isinstance(item, str):
            pdf.list_bullet(item)
        elif isinstance(item, dict):
            text = item.get("name", "") or item.get("title", "") or str(item)
            pdf.list_bullet(text)
    pdf.ln(2)


# Map section types to their renderers
_SECTION_RENDERERS = {
    "text": _render_section_text,
    "skills": _render_section_skills,
    "experience": _render_section_experience,
    "education": _render_section_education,
    "projects": _render_section_projects,
    "simple_list": _render_section_simple_list,
}


def generate_pdf(data: dict, template: str = "modern_green") -> bytes:
    pdf = ResumePDF(template=template)
    pdf.add_page()
    pdf.set_margins(18, 15, 18)
    pdf.set_y(15)

    name_align = pdf.cfg.get("name_align", "L")
    contact_align = pdf.cfg.get("contact_align", "L")
    name_size = pdf.cfg.get("name_size", 24)
    name_weight = "B" if pdf.cfg.get("name_bold", True) else ""

    # Name
    pdf.set_font("Helvetica", name_weight, name_size)
    pdf.set_text_color(*pdf.DARK)
    pdf.cell(0, 12, _clean_text(data.get("name", "")), align=name_align, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Contact
    contact = _clean_text(data.get("contact", ""))
    if contact:
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(*pdf.GRAY)
        contact_parts = [p.strip() for p in contact.split("|") if p.strip()]
        if pdf.cfg.get("contact_stacked"):
            # Each contact part on its own line, no separator (mirrors the minimal_clean layout).
            for part in contact_parts:
                pdf.cell(0, 5, part, align=contact_align, new_x="LMARGIN", new_y="NEXT")
        else:
            url_parts = [p for p in contact_parts if "http" in p or "linkedin" in p.lower() or "github" in p.lower()]
            non_url_parts = [p for p in contact_parts if p not in url_parts]
            if non_url_parts:
                pdf.cell(0, 5, " | ".join(non_url_parts), align=contact_align, new_x="LMARGIN", new_y="NEXT")
            for url in url_parts:
                pdf.cell(0, 5, url, align=contact_align, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Dynamic sections
    if "sections" in data:
        for section in data["sections"]:
            sec_type = section.get("type", "simple_list")
            renderer = _SECTION_RENDERERS.get(sec_type, _render_section_simple_list)
            renderer(pdf, section)
    else:
        # Backward compatibility: handle old fixed-schema format
        if data.get("summary"):
            pdf.section_header("Summary")
            pdf.body_text(data["summary"])
            pdf.ln(3)

        if data.get("skills"):
            pdf.section_header("Skills")
            for skill in data["skills"]:
                cat = _clean_text(skill.get("category", ""))
                items = _clean_text(skill.get("items", ""))
                pdf.set_font("Helvetica", "B", 9.5)
                pdf.set_text_color(*pdf.BODY)
                label = f"{cat}: "
                pdf.cell(pdf.get_string_width(label) + 1, 5, label)
                pdf.set_font("Helvetica", "", 9.5)
                remaining = pdf.w - pdf.r_margin - pdf.get_x()
                pdf.multi_cell(remaining, 4.5, items)
                pdf.ln(1)
            pdf.ln(2)

        if data.get("experience"):
            pdf.section_header("Experience")
            for job in data["experience"]:
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(*pdf.DARK)
                pdf.cell(0, 6, _clean_text(job.get("job_title", "")), new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*pdf.ACCENT)
                company_ctx = _clean_text(f"{job.get('company', '')} -- {job.get('context', '')}")
                pdf.cell(pdf.get_string_width(company_ctx) + 2, 5, company_ctx)
                pdf.set_text_color(*pdf.GRAY)
                meta = _clean_text(f"  |  {job.get('dates', '')}  |  {job.get('location', '')}")
                pdf.cell(0, 5, meta, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
                for b in job.get("bullets", []):
                    pdf.bullet(b)
                pdf.ln(3)

        edu = data.get("education", {})
        if edu:
            pdf.section_header("Education")
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(*pdf.BODY)
            pdf.cell(0, 6, _clean_text(edu.get("degree", "")), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*pdf.GRAY)
            edu_meta = _clean_text(f"{edu.get('school', '')}  |  {edu.get('dates', '')}  |  {edu.get('location', '')}")
            pdf.cell(0, 5, edu_meta)

    return pdf.output()


# ─────────────────────────────────────────────
# JD scraping helpers
# ─────────────────────────────────────────────

PLATFORM_DOMAINS = {
    "linkedin.com": "LinkedIn",
    "indeed.com": "Indeed",
    "glassdoor.com": "Glassdoor",
    "ziprecruiter.com": "ZipRecruiter",
    "monster.com": "Monster",
    "dice.com": "Dice",
    "lever.co": "Lever",
    "greenhouse.io": "Greenhouse",
    "workday.com": "Workday",
    "myworkdayjobs.com": "Workday",
    "smartrecruiters.com": "SmartRecruiters",
    "angel.co": "AngelList",
    "wellfound.com": "Wellfound",
    "builtin.com": "Built In",
    "simplyhired.com": "SimplyHired",
    "careerbuilder.com": "CareerBuilder",
    "welcometothejungle.com": "Welcome to the Jungle",
}

# All known platform names (lowercase) for filtering from titles
_PLATFORM_NAMES_LOWER = {v.lower() for v in PLATFORM_DOMAINS.values()} | {
    "linkedin", "indeed", "indeed.com", "glassdoor", "glassdoor.com",
    "otta", "welcome to the jungle",
}


def _is_platform_part(text):
    """Check if a title part is a platform name (not a real company)."""
    t = text.lower().strip()
    for pn in _PLATFORM_NAMES_LOWER:
        if pn in t:
            return True
    return False


def extract_job_metadata(url, title, text):
    """Detect platform from domain and parse company/position from page title."""
    parsed = urlparse(url)
    domain = parsed.netloc.lower().replace("www.", "")

    # Detect platform
    platform = "Other"
    for domain_key, platform_name in PLATFORM_DOMAINS.items():
        if domain_key in domain:
            platform = platform_name
            break

    company = ""
    position = ""

    if title:
        # LinkedIn: "Job Title at Company | LinkedIn"
        # Indeed: "Job Title - Company - Location | Indeed.com"
        # Glassdoor: "Company hiring Job Title in Location | Glassdoor"
        # WTTJ: "Company - Role | Welcome to the Jungle"

        if platform == "LinkedIn":
            m = re.match(r"^(.+?)\s+at\s+(.+?)(?:\s*\||\s*[-–]|\s*$)", title)
            if m:
                position = m.group(1).strip()
                company = m.group(2).strip()
        elif platform == "Indeed":
            parts = re.split(r"\s*[-–]\s*", title)
            if len(parts) >= 2:
                position = parts[0].strip()
                company = parts[1].strip()
        elif platform == "Glassdoor":
            m = re.match(r"^(.+?)\s+hiring\s+(.+?)(?:\s+in\s+|\s*\|)", title)
            if m:
                company = m.group(1).strip()
                position = m.group(2).strip()
        elif platform == "Welcome to the Jungle":
            # "Company - Role | Welcome to the Jungle (formerly Otta)"
            # Split on | first, take everything before the platform name
            main = re.split(r"\s*\|\s*", title)[0]
            parts = re.split(r"\s*[-–]\s*", main, maxsplit=1)
            if len(parts) >= 2:
                company = parts[0].strip()
                position = parts[1].strip()
            elif len(parts) == 1:
                position = parts[0].strip()

        # Generic fallback
        if not company and not position:
            # Split on | first to separate "content | platform branding"
            pipe_parts = re.split(r"\s*\|\s*", title)
            # Filter out parts that are platform names
            content_parts = [p.strip() for p in pipe_parts if not _is_platform_part(p)]
            if not content_parts:
                content_parts = [p.strip() for p in pipe_parts]

            # Use just the first content part (before any |), split on dash for company/role
            main = content_parts[0]
            sub = re.split(r"\s*[-–]\s*", main, maxsplit=1)
            if len(sub) >= 2:
                company = sub[0].strip()
                position = sub[1].strip()
            elif len(sub) == 1:
                position = sub[0].strip()

    return {
        "platform": platform,
        "company": company,
        "position": position,
        "url": url,
    }


# ─────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/scrape-jd", methods=["POST"])
def api_scrape_jd():
    """Scrape a job description from a URL using Apify website-content-crawler."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided."}), 400

        apify_token = data.get("apify_token", "").strip()
        url = data.get("url", "").strip()

        if not apify_token:
            return jsonify({"error": "Apify API token is required."}), 400
        if not url:
            return jsonify({"error": "Job URL is required."}), 400

        # Call Apify website-content-crawler actor
        actor_id = "apify~website-content-crawler"
        run_url = f"https://api.apify.com/v2/acts/{actor_id}/runs"

        run_input = {
            "startUrls": [{"url": url}],
            "maxCrawlPages": 1,
            "crawlerType": "playwright:firefox",
            "maxConcurrency": 1,
            "proxyConfiguration": {"useApifyProxy": True},
        }

        # Start the actor run synchronously (wait for finish)
        resp = requests.post(
            run_url,
            params={"token": apify_token, "waitForFinish": 120},
            json=run_input,
            timeout=180,
        )

        if resp.status_code == 401:
            return jsonify({"error": "Invalid Apify API token."}), 401
        if not resp.ok:
            return jsonify({"error": f"Apify API error: {resp.status_code} {resp.text[:200]}"}), 502

        run_data = resp.json().get("data", {})
        run_status = run_data.get("status")

        if run_status != "SUCCEEDED":
            return jsonify({"error": f"Apify crawl did not succeed (status: {run_status}). Try again."}), 502

        # Fetch results from the default dataset
        dataset_id = run_data.get("defaultDatasetId")
        if not dataset_id:
            return jsonify({"error": "No dataset returned from Apify."}), 502

        items_url = f"https://api.apify.com/v2/datasets/{dataset_id}/items"
        items_resp = requests.get(
            items_url,
            params={"token": apify_token, "format": "json"},
            timeout=30,
        )

        if not items_resp.ok:
            return jsonify({"error": "Failed to fetch crawl results."}), 502

        items = items_resp.json()
        if not items:
            return jsonify({"error": "No content was extracted from the URL. The page may require login."}), 404

        item = items[0]
        page_text = item.get("text", "")
        page_title = item.get("metadata", {}).get("title", "") or item.get("title", "")

        if not page_text:
            return jsonify({"error": "Page was loaded but no text content was found."}), 404

        metadata = extract_job_metadata(url, page_title, page_text)

        return jsonify({
            "success": True,
            "text": page_text,
            "title": page_title,
            "metadata": metadata,
        })

    except requests.Timeout:
        return jsonify({"error": "Apify request timed out. Try again."}), 504
    except Exception as e:
        return jsonify({"error": f"Scraping failed: {str(e)}"}), 500


@app.route("/api/jd-keywords", methods=["POST"])
def api_jd_keywords():
    """Preview the keywords the ATS scorer will look for, before the user spends tokens
    on a full tailoring run.

    Mode:
      - "heuristic" (default, free, no AI call) — regex-based extraction.
      - "ai" — same definition as the ATS scorer; uses tokens.
    """
    try:
        data = request.get_json() or {}
        jd_text = (data.get("jd") or "").strip()
        mode = (data.get("mode") or "heuristic").lower()
        if not jd_text:
            return jsonify({"error": "Job description is required."}), 400

        if mode == "ai":
            provider = (data.get("provider") or "anthropic").strip().lower()
            if provider == "claude":
                provider = "anthropic"
            api_key = (data.get("api_key") or "").strip()
            model = (data.get("model") or "").strip() or None
            base_url = (data.get("base_url") or "").strip() or None
            is_local = provider == "ollama" or (provider == "openai_compatible" and base_url and "localhost" in base_url)
            if not api_key and not is_local:
                return jsonify({"error": "API key is required for AI keyword extraction."}), 400
            keywords = extract_jd_keywords_ai(provider, api_key, jd_text, model=model, base_url=base_url)
            return jsonify({"success": True, "mode": "ai", "keywords": keywords})

        # heuristic mode
        keywords = extract_jd_keywords_heuristic(jd_text)
        return jsonify({"success": True, "mode": "heuristic", "keywords": keywords})

    except ValueError as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _provider_error_response(e)


@app.route("/api/tailor", methods=["POST"])
def api_tailor():
    try:
        provider = request.form.get("provider", "anthropic").strip().lower()
        if provider == "claude":
            provider = "anthropic"
        api_key = request.form.get("api_key", "").strip()
        model = request.form.get("model", "").strip() or None
        base_url = request.form.get("base_url", "").strip() or None
        fast_mode = request.form.get("fast_mode", "").strip().lower() in ("1", "true", "yes", "on")
        # Local providers (Ollama, openai_compatible pointing at localhost) often don't need a key.
        is_local = provider == "ollama" or (provider == "openai_compatible" and base_url and "localhost" in base_url)
        if not api_key and not is_local:
            return jsonify({"error": f"{provider.title()} API key is required."}), 400

        jd_text = request.form.get("jd", "").strip()
        resume_text_direct = request.form.get("resume_text", "").strip()

        # Parse prompt from file or text
        prompt_text = request.form.get("prompt", "").strip()
        if "prompt_file" in request.files:
            pf = request.files["prompt_file"]
            if pf.filename:
                pext = os.path.splitext(pf.filename)[1].lower()
                if pext not in PARSERS:
                    return jsonify({"error": f"Unsupported prompt file type: {pext}. Use .txt, .docx, or .pdf"}), 400
                prompt_text = PARSERS[pext](pf.read())

        if not prompt_text:
            prompt_text = DEFAULT_TAILORING_PROMPT
        if not jd_text:
            return jsonify({"error": "Job description is required."}), 400

        # Parse resume from file or text
        resume_text = ""
        personal_info = None
        if "resume_file" in request.files:
            f = request.files["resume_file"]
            if f.filename:
                ext = os.path.splitext(f.filename)[1].lower()
                if ext not in PARSERS:
                    return jsonify({"error": f"Unsupported file type: {ext}. Use .docx, .pdf, or .txt"}), 400
                file_bytes = f.read()
                resume_text = PARSERS[ext](file_bytes)
                # Extract original personal info from DOCX to preserve it
                if ext == ".docx":
                    personal_info = extract_personal_info_docx(file_bytes)

        if not resume_text and resume_text_direct:
            resume_text = resume_text_direct

        if not resume_text:
            return jsonify({"error": "Please upload a resume file or paste resume text."}), 400

        # Extract personal info from text for all file types (fallback)
        if not personal_info:
            personal_info = extract_personal_info_text(resume_text)

        # Detect sections from the original resume to preserve structure
        detected_sections = detect_resume_sections(resume_text)

        # Extract original title from resume text (first few lines, look for a title-like line)
        original_title = _extract_original_title(resume_text)

        result, ats = tailor_with_ats_target(
            api_key, resume_text, prompt_text, jd_text, provider, detected_sections,
            model=model, base_url=base_url, target=ATS_TARGET_SCORE, fast_mode=fast_mode,
        )

        # Force-override personal info with original values (never let AI change these)
        if personal_info:
            if personal_info.get("name"):
                result["name"] = personal_info["name"]
            if personal_info.get("contact"):
                result["contact"] = personal_info["contact"]

        # Force-override title with original (never let AI change the title)
        if original_title:
            result["title"] = original_title

        return jsonify({"success": True, "data": result, "ats": ats})

    except ValueError as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _provider_error_response(e)


@app.route("/api/answer-questions", methods=["POST"])
def api_answer_questions():
    """Answer job application questions using JD + resume context."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided."}), 400

        provider = data.get("provider", "anthropic").strip().lower()
        if provider == "claude":
            provider = "anthropic"
        api_key = data.get("api_key", "").strip()
        model = (data.get("model", "") or "").strip() or None
        base_url = (data.get("base_url", "") or "").strip() or None
        questions = data.get("questions", "").strip()
        jd_text = data.get("jd", "").strip()
        resume_json = data.get("resume", {})

        is_local = provider == "ollama" or (provider == "openai_compatible" and base_url and "localhost" in base_url)
        if not api_key and not is_local:
            return jsonify({"error": "API key is required."}), 400
        if not questions:
            return jsonify({"error": "Please paste at least one question."}), 400
        if not jd_text:
            return jsonify({"error": "Job description context is missing."}), 400

        # Rebuild a readable resume summary from the structured data
        resume_lines = []
        resume_lines.append(f"Name: {resume_json.get('name', '')}")
        resume_lines.append(f"Title: {resume_json.get('title', '')}")

        if "sections" in resume_json:
            # New dynamic sections format
            for section in resume_json["sections"]:
                heading = section.get("heading", "")
                sec_type = section.get("type", "")
                resume_lines.append(f"\n{heading}:")
                if sec_type == "text":
                    resume_lines.append(section.get("content", ""))
                elif sec_type == "skills":
                    for s in section.get("items", []):
                        resume_lines.append(f"  {s.get('category', '')}: {s.get('items', '')}")
                elif sec_type == "experience":
                    for job in section.get("items", []):
                        resume_lines.append(f"\n  {job.get('job_title', '')} at {job.get('company', '')} ({job.get('dates', '')})")
                        for b in job.get("bullets", []):
                            resume_lines.append(f"    - {b}")
                elif sec_type == "education":
                    items = section.get("items", [])
                    if isinstance(items, dict):
                        items = [items]
                    for edu in items:
                        resume_lines.append(f"  {edu.get('degree', '')} — {edu.get('school', '')}")
                elif sec_type == "projects":
                    for proj in section.get("items", []):
                        resume_lines.append(f"\n  {proj.get('name', '')}")
                        for b in proj.get("bullets", []):
                            resume_lines.append(f"    - {b}")
                else:
                    for item in section.get("items", []):
                        if isinstance(item, str):
                            resume_lines.append(f"  - {item}")
                        elif isinstance(item, dict):
                            resume_lines.append(f"  - {item.get('name', '') or item.get('title', '')}")
        else:
            # Old fixed-schema format (backward compat)
            resume_lines.append(f"\nSummary:\n{resume_json.get('summary', '')}")
            resume_lines.append("\nSkills:")
            for s in resume_json.get("skills", []):
                resume_lines.append(f"  {s.get('category', '')}: {s.get('items', '')}")
            resume_lines.append("\nExperience:")
            for job in resume_json.get("experience", []):
                resume_lines.append(f"\n  {job.get('job_title', '')} at {job.get('company', '')} ({job.get('dates', '')})")
                for b in job.get("bullets", []):
                    resume_lines.append(f"    - {b}")
            edu = resume_json.get("education", {})
            if edu:
                resume_lines.append(f"\nEducation: {edu.get('degree', '')} — {edu.get('school', '')}")
        resume_text = "\n".join(resume_lines)

        system_prompt = """You are an expert job application assistant. You help candidates write compelling,
authentic answers to job application questions. You have deep context about:
1. The candidate's background (their tailored resume)
2. The specific job they are applying for (the job description)

Rules:
- Write answers in FIRST PERSON as the candidate
- Keep answers concise but substantive (3-6 sentences per question unless it clearly needs more)
- Ground every answer in REAL experience from the resume — never fabricate
- Mirror the tone and keywords from the job description naturally
- Show enthusiasm for the specific role and company
- If a question asks about something not covered in the resume, craft an honest answer
  that pivots to relevant strengths rather than making things up
- For salary questions, suggest the candidate research market rates rather than giving a number
- For "why this company" questions, reference specific things from the JD that align with the candidate's experience"""

        user_msg = f"""Here is the candidate's resume:

{resume_text}

Here is the job description they are applying to:

{jd_text}

Please answer each of the following application questions. Format your response as a JSON array where each
element has "question" (the original question) and "answer" (your crafted response).

Questions:
{questions}

Return ONLY a JSON array, no markdown code fences, no extra text. Example format:
[{{"question": "Why do you want this role?", "answer": "Your answer here..."}}]"""

        full_prompt = f"{system_prompt}\n\n{user_msg}"
        # Note: Q&A returns a JSON array, but json_mode (response_format=json_object) on OpenAI
        # requires an object at the top level. Wrap accordingly: we ask for {"answers": [...]} when json_mode.
        is_openai_compat = provider in OPENAI_COMPATIBLE_BASE_URLS or provider == "openai_compatible"
        if is_openai_compat:
            full_prompt = full_prompt.replace(
                'Return ONLY a JSON array, no markdown code fences, no extra text. Example format:\n[{"question": "Why do you want this role?", "answer": "Your answer here..."}]',
                'Return ONLY a JSON object with an "answers" key. Example format:\n{"answers": [{"question": "Why do you want this role?", "answer": "Your answer here..."}]}',
            )
        raw = call_ai(provider, api_key, full_prompt, max_tokens=4000, model=model, base_url=base_url, json_mode=True)
        raw = _strip_code_fences(raw)

        try:
            parsed = _safe_json_loads(raw)
        except json.JSONDecodeError as e:
            return jsonify({"error": f"AI returned invalid JSON even after repair. First 500 chars: {raw[:500]!r}"}), 502
        # Accept either a bare array or {"answers": [...]}
        answers = parsed["answers"] if isinstance(parsed, dict) and "answers" in parsed else parsed
        if not isinstance(answers, list):
            return jsonify({"error": "AI response was not a list of answers."}), 502
        return jsonify({"success": True, "answers": answers})

    except ValueError as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _provider_error_response(e)


@app.route("/api/download-pdf", methods=["POST"])
def api_download_pdf():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No resume data provided."}), 400

        template = data.pop("template", "modern_green")
        pdf_bytes = generate_pdf(data, template=template)
        name_slug = data.get("name", "resume").replace(" ", "_")
        filename = f"{name_slug}.pdf"

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500



if __name__ == "__main__":
    debug = os.environ.get("FLASK_DEBUG", "").lower() in ("1", "true", "yes")
    app.run(debug=debug, port=5000)
